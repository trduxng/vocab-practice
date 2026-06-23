import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
doc = Document('document/vc_can_lam.docx')

# Show a few key paragraphs and tables to compare
print("=== DOCX STRUCTURE ===")
for i, para in enumerate(doc.paragraphs[:10]):
    if para.text.strip():
        print(f"P{i}: style='{para.style.name}' text='{para.text.strip()[:100]}'")

print("\n=== TABLES ===")
for i, table in enumerate(doc.tables[:3]):
    print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    for j, row in enumerate(table.rows):
        cells_text = [cell.text.strip()[:30] for cell in row.cells]
        print(f"  Row {j}: {cells_text}")

# Check images
print("\n=== IMAGES ===")
from docx.opc.constants import RELATIONSHIP_TYPE as RT
rels = doc.part.rels
img_count = 0
for rel in rels.values():
    if "image" in rel.reltype:
        img_count += 1
print(f"Images found: {img_count}")

# Check formatting
print("\n=== FORMATTING ===")
bold_count = 0
italic_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.bold: bold_count += 1
        if run.italic: italic_count += 1
print(f"Bold runs: {bold_count}, Italic runs: {italic_count}")
