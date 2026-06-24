from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("artifacts/design-doc/phuc-report-edit/Detail Design_Phuc_module_flow_updated.docx")
OUT = Path("artifacts/design-doc/phuc-report-edit/Detail Design_Phuc_module_controller_service_updated.docx")


def set_run_font(run, size=9, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.02


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="000000", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_text(cell, text, size=8.4, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, font="Times New Roman"):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, name=font)
    return p


def add_code(cell, text, size=7.0):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run(text.strip())
    set_run_font(r, size=size, name="Courier New")


def add_label(cell, label, text, size=8.2):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.02
    r1 = p.add_run(label)
    set_run_font(r1, size=size, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=size)


def move_paragraph_before(doc, ref, text="", style=None, bold=False):
    paragraph = doc.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    ref.addprevious(paragraph._p)
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=10 if style == "Heading 3" else 9, bold=bold or style == "Heading 3")
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def move_table_before(doc, ref, table):
    table._tbl.getparent().remove(table._tbl)
    ref.addprevious(table._tbl)


def build_summary_table(doc, module):
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = [Inches(1.55), Inches(5.10)]
    rows = [
        ("Mục đích", module["purpose"]),
        ("Thành phần", module["components"]),
        ("Luồng module", module["module_flow"]),
    ]
    for r, (label, value) in enumerate(rows):
        for c, cell in enumerate(table.rows[r].cells):
            clear_cell(cell)
            set_cell_width(cell, widths[c])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c == 0:
                shade_cell(cell, "D9EAF7")
                add_text(cell, label, size=8.7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_text(cell, value, size=8.3)
    return table


def build_function_table(doc, title, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = [Inches(1.95), Inches(2.10), Inches(2.60)]
    headers = [title, "Luồng xử lý", "Giả mã"]
    for i, cell in enumerate(table.rows[0].cells):
        clear_cell(cell)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        shade_cell(cell, "EAF2F8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(cell, headers[i], size=8.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for item in rows:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            clear_cell(cell)
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=95, start=95, bottom=95, end=95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        display_name = item["name"].replace("/", "\n")
        add_text(row.cells[0], display_name, size=8.0, bold=True)
        if item.get("file"):
            add_text(row.cells[0], item["file"], size=7.2, italic=True)
        if item.get("purpose"):
            add_label(row.cells[0], "Mục đích: ", item["purpose"], size=7.8)
        add_code(row.cells[1], item["flow"], size=6.9)
        add_code(row.cells[2], item["pseudo"], size=6.9)
    return table


MODULES = [
    {
        "name": "Module Auth - Xác thực và phân quyền",
        "purpose": "Xử lý đăng ký, đăng nhập, phát hành JWT và kiểm soát quyền truy cập API.",
        "components": "auth.routes.js; AuthController; AuthService; validate middleware; auth middleware; bảng Users, Roles, Permissions.",
        "module_flow": "Client -> auth.routes -> validate(schema) -> AuthController -> AuthService -> SQL Server -> response. Với API nội bộ: Client -> verifyToken -> checkPermission/checkAnyPermission -> Controller.",
        "controllers": [
            {
                "name": "AuthController.register(req,res,next)",
                "file": "controllers/auth.controller.js",
                "purpose": "Nhận request đăng ký và trả HTTP 201.",
                "flow": "Client POST /api/auth/register\n-> validate schemas.register\n-> controller lấy fullName,email,password\n-> gọi AuthService.register\n-> trả user mới\n-> lỗi chuyển next(error)",
                "pseudo": "FUNCTION register(req,res,next)\n  TRY\n    Lấy fullName,email,password từ req.body\n    user = AuthService.register(fullName,email,password)\n    res.status(201).json(user)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "AuthController.login(req,res,next)",
                "file": "controllers/auth.controller.js",
                "purpose": "Nhận request đăng nhập và trả token.",
                "flow": "Client POST /api/auth/login\n-> validate schemas.login\n-> controller lấy email,password\n-> gọi AuthService.login\n-> trả token,user,gamification nếu có",
                "pseudo": "FUNCTION login(req,res,next)\n  TRY\n    Lấy email,password từ req.body\n    result = AuthService.login(email,password)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "AuthService.register(fullName,email,password)",
                "file": "services/auth.service.js",
                "purpose": "Tạo user mới với role Learner.",
                "flow": "Service nhận dữ liệu từ controller\n-> query kiểm tra email\n-> hash password bằng bcrypt\n-> lấy RoleID Learner\n-> INSERT Users\n-> trả user object không chứa passwordHash",
                "pseudo": "FUNCTION register(fullName,email,password)\n  pool = mở kết nối CSDL\n  existing = SELECT Users WHERE Email=email\n  IF existing THEN throw 400\n  passwordHash = bcrypt.hash(password)\n  role = SELECT RoleID WHERE RoleName='Learner'\n  INSERT Users(fullName,email,passwordHash,role)\n  RETURN user mới\nEND FUNCTION",
            },
            {
                "name": "AuthService.login(email,password)",
                "file": "services/auth.service.js",
                "purpose": "Xác thực tài khoản và sinh JWT.",
                "flow": "Service query user active\n-> so sánh mật khẩu\n-> query permissions\n-> jwt.sign payload\n-> nếu Learner thì thưởng daily login\n-> trả token + user",
                "pseudo": "FUNCTION login(email,password)\n  user = SELECT active user by email\n  IF user không tồn tại THEN throw 401\n  valid = bcrypt.compare(password,user.PasswordHash)\n  IF valid=false THEN throw 401\n  permissions = query permission theo RoleID\n  token = jwt.sign({id,role,permissions})\n  IF role='Learner' THEN awardDailyLogin(userId)\n  RETURN {token,user}\nEND FUNCTION",
            },
            {
                "name": "verifyToken/checkPermission",
                "file": "middlewares/auth.js",
                "purpose": "Bảo vệ các module sau đăng nhập.",
                "flow": "Route gọi middleware\n-> verifyToken đọc Bearer token\n-> gán req.user\n-> checkPermission kiểm tra permissions\n-> hợp lệ thì next()",
                "pseudo": "FUNCTION verifyToken(req,res,next)\n  token = lấy Bearer token\n  IF thiếu token THEN return 401\n  payload = jwt.verify(token)\n  req.user = payload\n  next()\nEND FUNCTION\n\nFUNCTION checkPermission(permission)\n  IF req.user.permissions không chứa permission\n    return 403\n  next()\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Categories - Danh mục và chủ đề học",
        "purpose": "Cung cấp loại từ, topic học và tiến độ topic cho frontend.",
        "components": "categories.routes.js; CategoriesController; CategoriesService; bảng PartOfSpeeches, Topics, WordTopics, UserWordProgress.",
        "module_flow": "Client -> categories.routes -> CategoriesController -> CategoriesService -> SQL Server -> JSON danh mục/chủ đề.",
        "controllers": [
            {
                "name": "CategoriesController.getPartOfSpeeches",
                "file": "controllers/categories.controller.js",
                "purpose": "Trả danh sách loại từ.",
                "flow": "GET /part-of-speeches\n-> controller gọi service\n-> nhận recordset\n-> trả HTTP 200",
                "pseudo": "FUNCTION getPartOfSpeeches(req,res,next)\n  TRY\n    data = CategoriesService.getPartOfSpeeches()\n    res.status(200).json(data)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "CategoriesController.getTopics",
                "file": "controllers/categories.controller.js",
                "purpose": "Trả danh sách topic học.",
                "flow": "GET /topics\n-> lấy userId nếu có\n-> gọi CategoriesService.getTopics\n-> trả topics kèm progress",
                "pseudo": "FUNCTION getTopics(req,res,next)\n  TRY\n    userId = req.user?.id hoặc null\n    topics = CategoriesService.getTopics(userId)\n    res.status(200).json(topics)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "CategoriesService.getPartOfSpeeches",
                "file": "services/categories.service.js",
                "purpose": "Query bảng loại từ.",
                "flow": "Kết nối DB\n-> SELECT PartOfSpeeches\n-> ORDER BY tên\n-> return recordset",
                "pseudo": "FUNCTION getPartOfSpeeches()\n  pool = mở kết nối CSDL\n  result = SELECT * FROM PartOfSpeeches ORDER BY name\n  RETURN result.recordset\nEND FUNCTION",
            },
            {
                "name": "CategoriesService.getTopics(userId)",
                "file": "services/categories.service.js",
                "purpose": "Tính topic kèm wordCount/progress.",
                "flow": "Query Topics Published\n-> JOIN WordTopics/Words\n-> nếu có userId thì JOIN UserWordProgress\n-> tính learned/mastered/due/progress\n-> return topics",
                "pseudo": "FUNCTION getTopics(userId)\n  Query danh sách Topics Published\n  FOR mỗi topic\n    wordCount = đếm Words thuộc topic\n    IF userId tồn tại THEN\n      learnedCount = đếm progress của user\n      masteredCount = đếm từ mastered\n      dueCount = đếm từ đến hạn ôn\n      progressPercent = learnedCount/wordCount\n    END IF\n  END FOR\n  RETURN topics\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module User Learning - Học từ vựng, SRS, mini-test, notebook",
        "purpose": "Nghiệp vụ chính của Learner: học flashcard, ôn SRS, làm mini-test, xem tiến độ và quản lý notebook.",
        "components": "user.routes.js; UserController; UserService; GamificationService; bảng Words, Questions, UserWordProgress, ExerciseAttempts, MiniTestAttempts, UserVocabularyNotebook.",
        "module_flow": "Learner -> user.routes -> verifyToken -> UserController -> UserService -> SQL Server. Khi hoàn thành hoạt động học, UserService cập nhật progress và gọi GamificationService cộng XP.",
        "controllers": [
            {
                "name": "UserController.getFlashcards",
                "file": "controllers/user.controller.js",
                "purpose": "Nhận yêu cầu lấy flashcards.",
                "flow": "GET /flashcards\n-> lấy userId, topicId, mode\n-> gọi UserService.getDueFlashcards\n-> trả danh sách thẻ học",
                "pseudo": "FUNCTION getFlashcards(req,res,next)\n  TRY\n    userId = req.user.id\n    options = {topicId:req.query.topicId, mode:req.query.mode}\n    flashcards = UserService.getDueFlashcards(userId,options)\n    res.status(200).json(flashcards)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "UserController.submitAnswer",
                "file": "controllers/user.controller.js",
                "purpose": "Nhận đáp án luyện tập.",
                "flow": "POST /submit-answer\n-> lấy body answer/rating\n-> gọi UserService.submitAnswer\n-> trả kết quả đúng/sai, XP, nextReviewDate",
                "pseudo": "FUNCTION submitAnswer(req,res,next)\n  TRY\n    payload = req.body + userId từ token\n    result = UserService.submitAnswer(payload)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "UserController.submitMiniTest",
                "file": "controllers/user.controller.js",
                "purpose": "Nhận bài làm mini-test.",
                "flow": "POST /minitests/:id/submit\n-> lấy testId và answers\n-> gọi UserService.submitMiniTestBatch\n-> trả điểm và chi tiết",
                "pseudo": "FUNCTION submitMiniTest(req,res,next)\n  TRY\n    testId = req.params.id\n    answers = req.body.answers\n    result = UserService.submitMiniTestBatch(userId,testId,answers)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "UserController.getProgressAnalytics",
                "file": "controllers/user.controller.js",
                "purpose": "Trả dữ liệu biểu đồ tiến độ.",
                "flow": "GET /progress/analytics\n-> lấy userId\n-> gọi UserService.getProgressAnalytics\n-> trả analytics",
                "pseudo": "FUNCTION getProgressAnalytics(req,res,next)\n  TRY\n    analytics = UserService.getProgressAnalytics(req.user.id)\n    res.status(200).json(analytics)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "Notebook controllers",
                "file": "controllers/user.controller.js",
                "purpose": "CRUD sổ tay từ vựng.",
                "flow": "GET/POST/PUT/DELETE /notebook\n-> controller đọc params/body\n-> gọi service tương ứng\n-> trả entry/list/success",
                "pseudo": "FUNCTION notebookController(req,res,next)\n  TRY\n    IF GET THEN getNotebook(userId,page,pageSize)\n    IF POST THEN addNotebookEntry(userId,wordId,note)\n    IF PUT THEN updateNotebookEntry(id,userId,data)\n    IF DELETE THEN deleteNotebookEntry(id,userId)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "UserService.getDueFlashcards",
                "file": "services/user.service.js",
                "purpose": "Chọn từ cần học/ôn theo SRS.",
                "flow": "Chuẩn hóa topic/mode\n-> lấy limit SRS của user\n-> query Words Published + UserWordProgress\n-> ưu tiên từ due/sai nhiều/từ mới\n-> trả flashcards",
                "pseudo": "FUNCTION getDueFlashcards(userId,options)\n  limit = lấy SRSReviewLimit của user\n  Query Words Published LEFT JOIN UserWordProgress\n  IF mode='new' THEN lọc từ chưa học\n  IF mode='daily' THEN lọc NextReviewDate <= now\n  ORDER BY độ ưu tiên SRS\n  RETURN flashcards\nEND FUNCTION",
            },
            {
                "name": "UserService.submitAnswer",
                "file": "services/user.service.js",
                "purpose": "Ghi attempt và cập nhật feedback SRS.",
                "flow": "Gọi stored procedure chấm câu hỏi\n-> lấy canonicalWordId\n-> nếu có reviewRating thì cập nhật EaseFactor/NextReviewDate\n-> nếu LearnWord thì awardXP\n-> trả result",
                "pseudo": "FUNCTION submitAnswer(payload)\n  result = EXEC usp_SubmitQuestionAttempt\n  wordId = result.WordID\n  IF wordId rỗng THEN throw error\n  IF reviewRating in Again/Hard/Good/Easy THEN\n    UPDATE UserWordProgress\n    SET EaseFactor, NextReviewDate theo rating\n  END IF\n  IF activityType='LearnWord' THEN awardXP\n  RETURN result + reviewFeedback + xp\nEND FUNCTION",
            },
            {
                "name": "UserService.submitWordReview",
                "file": "services/user.service.js",
                "purpose": "Cập nhật tiến độ một từ sau flashcard.",
                "flow": "MERGE UserWordProgress\n-> đúng tăng mastery, sai giảm mastery\n-> tính nextReviewDate theo rating\n-> cập nhật MemoryStatus\n-> cộng XP nếu cần",
                "pseudo": "FUNCTION submitWordReview(payload)\n  MERGE progress by userId,wordId\n  IF matched THEN UPDATE mastery/repetition/streak\n  ELSE INSERT progress mới\n  NextReviewDate = tính theo Again/Hard/Good/Easy\n  MemoryStatus = Learning/Reviewing/Mastered/Lapsed\n  RETURN progress mới\nEND FUNCTION",
            },
            {
                "name": "UserService.getSmartReviewQueue",
                "file": "services/user.service.js",
                "purpose": "Tạo hàng đợi ôn thông minh.",
                "flow": "Query progress có NextReviewDate trong 7 ngày\n-> tính priorityScore\n-> từ sai nhiều/quá hạn có điểm cao\n-> trả TOP limit",
                "pseudo": "FUNCTION getSmartReviewQueue(userId,limit)\n  limit = clamp(1,50)\n  SELECT Words JOIN UserWordProgress\n  WHERE NextReviewDate <= now + 7 days\n  priorityScore = overdueHours * wrongMultiplier\n  ORDER BY priorityScore DESC, MasteryLevel ASC\n  RETURN queue\nEND FUNCTION",
            },
            {
                "name": "UserService.submitMiniTestBatch",
                "file": "services/user.service.js",
                "purpose": "Chấm mini-test theo transaction.",
                "flow": "Begin transaction\n-> lấy câu hỏi của test\n-> validate answers\n-> loop từng answer: insert attempt, merge progress\n-> insert MiniTestAttempts\n-> commit, awardXP",
                "pseudo": "FUNCTION submitMiniTestBatch(userId,testId,answers)\n  BEGIN TRANSACTION\n  questions = SELECT questions by testId\n  IF answers.length != questions.length THEN throw error\n  FOR answer IN answers\n    isCorrect = compare(answer,CorrectAnswer)\n    INSERT ExerciseAttempts\n    MERGE UserWordProgress\n  END FOR\n  INSERT MiniTestAttempts(score)\n  COMMIT\n  awardXP(MiniTestComplete)\n  RETURN score/details\nEND FUNCTION",
            },
            {
                "name": "UserService.getProgressAnalytics",
                "file": "services/user.service.js",
                "purpose": "Tổng hợp dashboard học tập.",
                "flow": "Query progress, attempts, topic mastery\n-> tính accuracy, mastery distribution, activity timeline\n-> trả object cho chart",
                "pseudo": "FUNCTION getProgressAnalytics(userId)\n  learned = count UserWordProgress\n  mastered = count MemoryStatus='Mastered'\n  accuracy = avg ExerciseAttempts.IsCorrect\n  topicProgress = group by topic\n  timeline = group activity by date\n  RETURN analytics\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Learning Path & Progress",
        "purpose": "Sinh lộ trình học TOEIC và thống kê tiến độ tổng quan.",
        "components": "LearningPathController; LearningPathService; ProgressController; ProgressService; LearningPathLevels, LearningPathTopics, UserWordProgress.",
        "module_flow": "Client -> learning-path/progress routes -> Controller -> Service -> SQL Server -> roadmap/progress response.",
        "controllers": [
            {
                "name": "LearningPathController.getRoadmap",
                "file": "controllers/learning-path.controller.js",
                "purpose": "Trả roadmap cho learner.",
                "flow": "GET /learning-path\n-> lấy userId\n-> gọi LearningPathService.getRoadmap\n-> trả roadmap",
                "pseudo": "FUNCTION getRoadmap(req,res,next)\n  TRY\n    roadmap = LearningPathService.getRoadmap(req.user.id)\n    res.status(200).json(roadmap)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "ProgressController.getProgress/getStats",
                "file": "controllers/progress.controller.js",
                "purpose": "Trả tiến độ và thống kê.",
                "flow": "GET /progress hoặc /progress/stats\n-> lấy userId\n-> gọi ProgressService\n-> trả JSON",
                "pseudo": "FUNCTION progressController(req,res,next)\n  TRY\n    IF route='/' THEN result = ProgressService.getProgress(userId)\n    IF route='/stats' THEN result = ProgressService.getStats(userId)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "LearningPathService.getRoadmap",
                "file": "services/learning-path.service.js",
                "purpose": "Lấy roadmap theo level/topic.",
                "flow": "ensureSchema\n-> syncPublishedTopics\n-> query levels/topics\n-> query user progress\n-> buildRoadmap",
                "pseudo": "FUNCTION getRoadmap(userId)\n  ensureSchema()\n  syncPublishedTopics()\n  levels = SELECT LearningPathLevels\n  topics = SELECT LearningPathTopics JOIN Topics\n  progress = SELECT UserWordProgress by userId\n  RETURN buildRoadmap(levels,topics,progress)\nEND FUNCTION",
            },
            {
                "name": "LearningPathService.buildRoadmap",
                "file": "services/learning-path.service.js",
                "purpose": "Định dạng roadmap cho UI.",
                "flow": "Duyệt từng level\n-> gắn topics thuộc level\n-> tính progressPercent\n-> xác định trạng thái completed/available/locked",
                "pseudo": "FUNCTION buildRoadmap(levels,topics,progress)\n  FOR level IN levels\n    levelTopics = topics.filter(level)\n    FOR topic IN levelTopics\n      topic.progress = tính theo UserWordProgress\n      topic.status = completed/available/locked\n    END FOR\n  END FOR\n  RETURN roadmap\nEND FUNCTION",
            },
            {
                "name": "ProgressService.getProgress/getStats",
                "file": "services/progress.service.js",
                "purpose": "Tính tiến độ học.",
                "flow": "Query UserWordProgress/Attempts\n-> đếm learned/mastered/reviewing\n-> tính accuracy/streak\n-> trả summary",
                "pseudo": "FUNCTION getProgress(userId)\n  progress = SELECT UserWordProgress WHERE userId\n  RETURN counts by MemoryStatus\nEND FUNCTION\n\nFUNCTION getStats(userId)\n  attempts = SELECT ExerciseAttempts WHERE userId\n  accuracy = avg isCorrect\n  streak = tính ngày học liên tiếp\n  RETURN stats\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Gamification - XP, Level, Achievement",
        "purpose": "Cộng XP, tính level/streak và mở khóa huy hiệu cho learner.",
        "components": "GamificationController; GamificationService; UserXPEvents; Achievements; UserAchievements.",
        "module_flow": "Hoạt động học -> Controller/UserService -> GamificationService -> ghi XP event -> cập nhật user -> kiểm tra achievement -> response.",
        "controllers": [
            {
                "name": "GamificationController.getProfile",
                "file": "controllers/gamification.controller.js",
                "purpose": "Trả hồ sơ gamification.",
                "flow": "GET /gamification/profile\n-> lấy userId\n-> gọi service.getProfile\n-> trả XP/level/achievements",
                "pseudo": "FUNCTION getProfile(req,res,next)\n  TRY\n    profile = GamificationService.getProfile(req.user.id)\n    res.status(200).json(profile)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "completePractice/markAchievementsSeen",
                "file": "controllers/gamification.controller.js",
                "purpose": "Cộng XP phiên luyện tập và đánh dấu huy hiệu đã xem.",
                "flow": "Practice complete -> awardXP\nAchievements seen -> markAchievementsSeen\nController trả success/result",
                "pseudo": "FUNCTION gamificationActions(req,res,next)\n  TRY\n    IF completePractice THEN awardXP(userId,PracticeComplete)\n    IF markSeen THEN markAchievementsSeen(userId,ids)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "GamificationService.awardXP",
                "file": "services/gamification.service.js",
                "purpose": "Cộng XP theo sự kiện.",
                "flow": "ensureSchema\n-> xác định XP_REWARDS\n-> chống cộng trùng bằng sourceKey\n-> insert UserXPEvents\n-> update Users.TotalXP\n-> checkAchievements",
                "pseudo": "FUNCTION awardXP(userId,event)\n  xp = XP_REWARDS[eventType]\n  IF sourceKey đã tồn tại THEN return xpGained=0\n  INSERT UserXPEvents\n  UPDATE Users SET TotalXP += xp\n  levelState = getLevelState(totalXP)\n  achievements = checkAchievements(userId)\n  RETURN xpGained,totalXP,levelState,achievements\nEND FUNCTION",
            },
            {
                "name": "GamificationService.checkAchievements",
                "file": "services/gamification.service.js",
                "purpose": "Mở khóa achievement.",
                "flow": "Lấy metrics\n-> duyệt achievements\n-> kiểm tra điều kiện\n-> insert UserAchievements nếu đạt",
                "pseudo": "FUNCTION checkAchievements(userId)\n  metrics = getMetrics(userId)\n  achievements = getAchievements()\n  FOR achievement IN achievements\n    IF user đạt điều kiện AND chưa unlock THEN\n      INSERT UserAchievements\n    END IF\n  END FOR\n  RETURN newlyUnlocked\nEND FUNCTION",
            },
            {
                "name": "GamificationService.getProfile",
                "file": "services/gamification.service.js",
                "purpose": "Tổng hợp profile.",
                "flow": "ensureSchema\n-> getMetrics\n-> getAchievements\n-> getLevelState\n-> trả profile",
                "pseudo": "FUNCTION getProfile(userId)\n  ensureSchema()\n  metrics = getMetrics(userId)\n  achievements = getAchievements(userId,metrics)\n  levelState = getLevelState(metrics.totalXP)\n  RETURN {metrics,achievements,levelState}\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Admin - Quản trị hệ thống và nội dung",
        "purpose": "Quản trị toàn bộ nội dung, tài khoản, kiểm duyệt, thống kê, báo cáo, audit log và thông báo.",
        "components": "admin.routes.js; AdminController; AdminService; ReportService; auth/validate middleware; các bảng Words, Topics, Questions, MiniTests, Users, AuditLogs, Reports, Notifications.",
        "module_flow": "Admin -> /api/admin -> verifyToken -> checkPermission/checkAnyPermission -> validate -> AdminController -> AdminService/ReportService -> SQL Server -> audit/review log -> response.",
        "controllers": [
            {
                "name": "AdminController.getTopics/createTopic/updateTopic/deleteTopic",
                "file": "controllers/admin.controller.js",
                "purpose": "Điều phối CRUD chủ đề.",
                "flow": "Route /topics\n-> kiểm tra quyền MANAGE_TOPICS/MANAGE_WORDS\n-> validate body nếu POST/PUT\n-> controller đọc params/query/body\n-> gọi AdminService",
                "pseudo": "FUNCTION adminTopicController(req,res,next)\n  TRY\n    IF GET THEN result = AdminService.getTopics(page,limit,filters)\n    IF POST THEN result = AdminService.createTopic(req.body,adminId)\n    IF PUT THEN result = AdminService.updateTopic(id,body,adminId)\n    IF DELETE THEN result = AdminService.deleteTopic(id,adminId)\n    res.status(successCode).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "AdminController.getWords/createWord/updateWord/deleteWord/import",
                "file": "controllers/admin.controller.js",
                "purpose": "Điều phối quản lý từ vựng.",
                "flow": "Route /words\n-> check MANAGE_WORDS\n-> controller đọc filter/body/csv text\n-> gọi service get/create/update/delete/preview/bulkImport\n-> trả JSON",
                "pseudo": "FUNCTION adminWordController(req,res,next)\n  TRY\n    SWITCH route\n      GET /words -> getWords(page,limit,filters)\n      POST /words -> createWord(body,adminId)\n      PUT /words/:id -> updateWord(id,body,adminId)\n      DELETE -> deleteWord(id,adminId)\n      import -> previewWordImport/bulkInsertWords(text)\n    END SWITCH\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "AdminController question/miniTest handlers",
                "file": "controllers/admin.controller.js",
                "purpose": "Điều phối câu hỏi và mini-test.",
                "flow": "Route questions/minitests\n-> check MANAGE_QUESTIONS/MANAGE_TESTS\n-> validate body\n-> gọi service tương ứng\n-> trả danh sách/chi tiết/status",
                "pseudo": "FUNCTION adminAssessmentController(req,res,next)\n  TRY\n    IF question route THEN call AdminService question method\n    IF minitest route THEN call AdminService miniTest method\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "AdminController user/review/report/notification handlers",
                "file": "controllers/admin.controller.js",
                "purpose": "Điều phối user, duyệt nội dung, báo cáo, thông báo.",
                "flow": "Route students/review/reports/notifications\n-> check quyền phù hợp\n-> controller lấy params/body\n-> gọi AdminService hoặc ReportService\n-> trả kết quả",
                "pseudo": "FUNCTION adminOperationController(req,res,next)\n  TRY\n    IF students THEN call user management service\n    IF content-review THEN approve/reject/archive/logs\n    IF reports THEN ReportService.get/update\n    IF notifications THEN sendAnnouncement/createDailyReminders\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "AdminService.getTopics/createTopic/updateTopic/deleteTopic",
                "file": "services/admin.service.js",
                "purpose": "Thực thi CRUD Topics.",
                "flow": "get: build filter + paginate\ncreate/update: transaction topic + metadata\ndelete: kiểm tra liên kết rồi xóa/ẩn\naudit log sau thay đổi",
                "pseudo": "FUNCTION adminTopicService(action)\n  IF get THEN query Topics + wordCount + miniTestCount\n  IF create THEN INSERT Topics\n  IF update THEN UPDATE Topics WHERE TopicID\n  IF delete THEN kiểm tra WordTopics/MiniTests rồi DELETE/Archive\n  logAdminAction(adminId,action,'Topic',id)\n  RETURN result\nEND FUNCTION",
            },
            {
                "name": "AdminService.getWords/createWord/updateWord/deleteWord",
                "file": "services/admin.service.js",
                "purpose": "Quản lý Words, topics, examples.",
                "flow": "get: query Words + relations\ncreate/update: transaction\n-> Words\n-> WordTopics\n-> ExampleSentences\ndelete/archive theo quyền",
                "pseudo": "FUNCTION adminWordService(action)\n  BEGIN TRANSACTION nếu create/update\n  IF create THEN INSERT Words\n  IF update THEN UPDATE Words\n  IF topicIds THEN replace WordTopics\n  IF examples THEN replace ExampleSentences\n  COMMIT\n  logAdminAction\n  RETURN word/result\nEND FUNCTION",
            },
            {
                "name": "AdminService.previewWordImport/bulkInsertWords",
                "file": "services/admin.service.js",
                "purpose": "Nhập hàng loạt từ vựng.",
                "flow": "preview parse csv -> validate rows\nbulk parse -> transaction -> insert words/topics/examples -> commit/rollback",
                "pseudo": "FUNCTION bulkInsertWords(input,adminId)\n  rows = parseDelimitedImport(input)\n  validRows,errors = validate rows\n  BEGIN TRANSACTION\n  FOR row IN validRows\n    INSERT Words\n    INSERT WordTopics\n    INSERT ExampleSentences\n  END FOR\n  COMMIT\n  RETURN {inserted,errors}\nEND FUNCTION",
            },
            {
                "name": "AdminService question/miniTest services",
                "file": "services/admin.service.js",
                "purpose": "Quản lý Questions và MiniTests.",
                "flow": "Questions: CRUD theo WordID\nMiniTests: CRUD test + items\npublish/archive cập nhật status và logs",
                "pseudo": "FUNCTION assessmentService(action)\n  IF question THEN validate WordID, answer, options; INSERT/UPDATE/DELETE Questions\n  IF miniTest THEN BEGIN TRANSACTION; INSERT/UPDATE MiniTests; replace MiniTestItems; COMMIT\n  IF publish/archive THEN setMiniTestStatus\n  RETURN result\nEND FUNCTION",
            },
            {
                "name": "AdminService user/content/report/notification services",
                "file": "services/admin.service.js; services/report.service.js",
                "purpose": "Vận hành hệ thống.",
                "flow": "Users: query/update role/status\nContent: update status + ContentReviewLogs\nReports: get/update status\nNotifications: insert announcement/reminder",
                "pseudo": "FUNCTION adminOperationService(action)\n  IF user THEN SELECT/UPDATE Users, Roles\n  IF content THEN UPDATE entity ContentStatus; INSERT ContentReviewLogs\n  IF report THEN ReportService.getReports/updateReport\n  IF notification THEN INSERT Notifications\n  logAdminAction\n  RETURN result\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Creator - Biên tập nội dung học",
        "purpose": "Cho Content Creator tạo topic, word, question, mini-test, media và gửi duyệt.",
        "components": "creator.routes.js; CreatorController; CreatorService; validate/auth middleware; Topics, Words, Questions, MiniTests, Media.",
        "module_flow": "Creator -> /api/creator -> verifyToken -> permission -> validate -> CreatorController -> CreatorService -> SQL Server. Nội dung tạo ở Draft, submit-review chuyển PendingReview.",
        "controllers": [
            {
                "name": "CreatorController dashboard handlers",
                "file": "controllers/creator.controller.js",
                "purpose": "Trả dashboard và analytics cho creator.",
                "flow": "GET dashboard/summary/analytics\n-> lấy userId\n-> gọi CreatorService stats/analytics\n-> trả dữ liệu",
                "pseudo": "FUNCTION creatorDashboardController(req,res,next)\n  TRY\n    userId = req.user.id\n    result = CreatorService.getDashboardStats/getContentSummary/getAnalytics(userId)\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "CreatorController topic/word/question handlers",
                "file": "controllers/creator.controller.js",
                "purpose": "Điều phối CRUD nội dung cơ bản.",
                "flow": "Route topics/words/questions\n-> check permission\n-> validate body\n-> controller gọi CreatorService\n-> trả entity/status",
                "pseudo": "FUNCTION creatorContentController(req,res,next)\n  TRY\n    IF topic route THEN call topic service\n    IF word route THEN call word service\n    IF question route THEN call question service\n    IF submit-review THEN call submitForReview wrapper\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "CreatorController miniTest/media handlers",
                "file": "controllers/creator.controller.js",
                "purpose": "Điều phối mini-test và media.",
                "flow": "Route mini-tests/media\n-> check permission\n-> validate\n-> gọi CreatorService\n-> trả result",
                "pseudo": "FUNCTION creatorExtraController(req,res,next)\n  TRY\n    IF miniTest route THEN call miniTest service\n    IF media route THEN call media service\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "CreatorService.create/update/deleteTopic",
                "file": "services/creator.service.js",
                "purpose": "Creator quản lý topic của mình.",
                "flow": "Kiểm tra ownership\n-> insert/update/delete Topics\n-> status mặc định Draft\n-> trả topic",
                "pseudo": "FUNCTION creatorTopicService(action)\n  IF create THEN INSERT Topics(ContentStatus='Draft')\n  IF update/delete THEN kiểm tra CreatedByUserID=userId\n  UPDATE/DELETE Topics\n  RETURN topic/status\nEND FUNCTION",
            },
            {
                "name": "CreatorService.createWord/updateWord/deleteWord",
                "file": "services/creator.service.js",
                "purpose": "Biên soạn từ vựng và ví dụ.",
                "flow": "Begin transaction\n-> kiểm tra partOfSpeech/topic\n-> kiểm tra trùng từ\n-> insert/update Words\n-> insert WordTopics/ExampleSentences\n-> commit",
                "pseudo": "FUNCTION creatorWordService(data,userId)\n  BEGIN TRANSACTION\n  Validate PartOfSpeechID\n  Validate topic ownership/Published\n  Check duplicate Term + PartOfSpeech\n  INSERT/UPDATE Words status Draft\n  Replace WordTopics and ExampleSentences\n  COMMIT\n  RETURN word\nEND FUNCTION",
            },
            {
                "name": "CreatorService.createQuestion/updateQuestion",
                "file": "services/creator.service.js",
                "purpose": "Tạo câu hỏi luyện tập.",
                "flow": "Kiểm tra word được phép dùng\n-> validate question/options\n-> insert/update Questions\n-> status Draft",
                "pseudo": "FUNCTION creatorQuestionService(data,userId)\n  Validate WordID belongs to creator or Published\n  Validate question text, correctAnswer, options\n  INSERT/UPDATE Questions(ContentStatus='Draft')\n  RETURN question\nEND FUNCTION",
            },
            {
                "name": "CreatorService.createMiniTest/updateMiniTest/items",
                "file": "services/creator.service.js",
                "purpose": "Tạo bài test từ câu hỏi.",
                "flow": "validateMiniTestReferences\n-> transaction MiniTests + MiniTestItems\n-> add/remove item kiểm tra ownership\n-> return miniTest",
                "pseudo": "FUNCTION creatorMiniTestService(data,userId)\n  BEGIN TRANSACTION\n  validateMiniTestReferences(topicId,questionIds,userId)\n  INSERT/UPDATE MiniTests(ContentStatus='Draft')\n  INSERT/DELETE MiniTestItems\n  COMMIT\n  RETURN miniTest\nEND FUNCTION",
            },
            {
                "name": "CreatorService.submitForReview",
                "file": "services/creator.service.js",
                "purpose": "Gửi nội dung sang admin duyệt.",
                "flow": "Kiểm tra entity thuộc creator\n-> status phải Draft/Rejected\n-> update PendingReview\n-> trả entity",
                "pseudo": "FUNCTION submitForReview(table,idColumn,id,userId)\n  entity = SELECT WHERE id AND CreatedByUserID=userId\n  IF không tồn tại THEN throw 404\n  IF status không thuộc Draft/Rejected THEN throw error\n  UPDATE entity SET ContentStatus='PendingReview'\n  RETURN updated entity\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Review - Kiểm duyệt nội dung",
        "purpose": "Reviewer/Admin duyệt, từ chối, lưu trữ và xem log nội dung.",
        "components": "review.routes.js; ReviewController; ReviewService; ContentReviewLogs; Topics, Words, Questions, MiniTests.",
        "module_flow": "Reviewer -> /api/review -> verifyToken -> REVIEW_CONTENT permission -> ReviewController -> ReviewService -> SQL transaction -> response.",
        "controllers": [
            {
                "name": "ReviewController.getPending",
                "file": "controllers/review.controller.js",
                "purpose": "Trả danh sách nội dung chờ duyệt.",
                "flow": "GET /pending\n-> check REVIEW_CONTENT\n-> gọi ReviewService.getPendingContent\n-> trả danh sách",
                "pseudo": "FUNCTION getPending(req,res,next)\n  TRY\n    data = ReviewService.getPendingContent()\n    res.status(200).json(data)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "ReviewController.approve/reject/archive/logs",
                "file": "controllers/review.controller.js",
                "purpose": "Điều phối hành động duyệt.",
                "flow": "POST/GET entityType/entityId\n-> lấy params + adminId\n-> gọi ReviewService\n-> trả success/logs",
                "pseudo": "FUNCTION reviewActionController(req,res,next)\n  TRY\n    entityType = req.params.entityType\n    entityId = req.params.entityId\n    IF approve THEN result = ReviewService.approve(entityType,entityId,adminId)\n    IF reject THEN result = ReviewService.reject(entityType,entityId,adminId,reason)\n    IF archive THEN result = ReviewService.archive(entityType,entityId,adminId)\n    IF logs THEN result = ReviewService.getReviewLogs(entityType,entityId)\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "ReviewService.getPendingContent",
                "file": "services/review.service.js",
                "purpose": "Gộp nội dung PendingReview.",
                "flow": "UNION Topics/Words/Questions/MiniTests\n-> JOIN Users creator\n-> order by createdAt\n-> return list",
                "pseudo": "FUNCTION getPendingContent()\n  SELECT PendingReview FROM Topics\n  UNION ALL SELECT PendingReview FROM Words\n  UNION ALL SELECT PendingReview FROM Questions\n  UNION ALL SELECT PendingReview FROM MiniTests\n  ORDER BY createdAt ASC\n  RETURN recordset\nEND FUNCTION",
            },
            {
                "name": "ReviewService.approve/reject/archive",
                "file": "services/review.service.js",
                "purpose": "Cập nhật trạng thái nội dung.",
                "flow": "Resolve entityType\n-> begin transaction\n-> update status\n-> update reviewed fields\n-> insert ContentReviewLogs\n-> commit",
                "pseudo": "FUNCTION reviewStatusAction(entityType,entityId,adminId,action)\n  e = _resolveEntity(entityType)\n  BEGIN TRANSACTION\n  UPDATE e.table SET ContentStatus = actionStatus\n  IF MiniTest THEN update IsPublished\n  INSERT ContentReviewLogs(oldStatus,newStatus,comment)\n  COMMIT\n  RETURN true\nEND FUNCTION",
            },
            {
                "name": "ReviewService.getReviewLogs",
                "file": "services/review.service.js",
                "purpose": "Lấy lịch sử kiểm duyệt.",
                "flow": "Resolve entity\n-> query ContentReviewLogs\n-> join Users actionBy\n-> order desc",
                "pseudo": "FUNCTION getReviewLogs(entityType,entityId)\n  e = _resolveEntity(entityType)\n  logs = SELECT ContentReviewLogs WHERE EntityType=e.type AND EntityID=entityId\n  JOIN Users reviewer\n  ORDER BY CreatedAt DESC\n  RETURN logs\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module AI - Hỗ trợ soạn nội dung TOEIC",
        "purpose": "Gợi ý nghĩa, ví dụ, bản dịch và dữ liệu từ điển cho nội dung TOEIC.",
        "components": "ai.routes.js; AiController; AiService; AI provider/dictionary.",
        "module_flow": "Admin/Creator -> /api/ai -> verifyToken/permission -> AiController -> AiService -> AI provider -> parse/normalize -> response.",
        "controllers": [
            {
                "name": "AiController.suggestWordContent",
                "file": "controllers/ai.controller.js",
                "purpose": "Nhận yêu cầu gợi ý từ vựng.",
                "flow": "POST /suggest-word-content\n-> validate quyền/body\n-> gọi AiService.suggestWordContent\n-> trả suggestion",
                "pseudo": "FUNCTION suggestWordContent(req,res,next)\n  TRY\n    payload = req.body\n    result = AiService.suggestWordContent(payload)\n    res.status(200).json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "AiService.suggestWordContent",
                "file": "services/ai.service.js",
                "purpose": "Sinh gợi ý đầy đủ cho một từ.",
                "flow": "Clean input\n-> tạo prompt TOEIC\n-> gọi AI withTimeout\n-> extract response text\n-> parse JSON\n-> normalize examples",
                "pseudo": "FUNCTION suggestWordContent(payload)\n  term = cleanText(payload.term)\n  prompt = build TOEIC prompt\n  response = call AI provider with timeout\n  text = extractResponseText(response)\n  data = parseJsonText(text)\n  data.examples = normalizeExamples(data.examples)\n  RETURN data\nEND FUNCTION",
            },
            {
                "name": "generateToeicExamples/translate/lookupDictionary",
                "file": "services/ai.service.js",
                "purpose": "Các hàm hỗ trợ AI.",
                "flow": "generate: prompt ví dụ TOEIC\ntranslate: prompt dịch tiếng Việt\nlookup: tra cứu dictionary\n-> normalize output",
                "pseudo": "FUNCTION aiHelpers(action,input)\n  IF generateExamples THEN call AI and normalize examples\n  IF translate THEN call AI translate and return Vietnamese text\n  IF lookupDictionary THEN query dictionary source\n  RETURN normalized result\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Report & Notification - Báo cáo và thông báo",
        "purpose": "Người học gửi báo cáo; admin xử lý báo cáo và gửi thông báo/lời nhắc học.",
        "components": "UserController.createReport; ReportService; AdminController report/notification handlers; AdminService notification methods; Notifications, Reports.",
        "module_flow": "Learner -> createReport -> ReportService -> DB. Admin -> reports/notifications routes -> AdminController -> ReportService/AdminService -> DB -> User đọc notifications.",
        "controllers": [
            {
                "name": "UserController.createReport",
                "file": "controllers/user.controller.js",
                "purpose": "Nhận báo cáo từ learner.",
                "flow": "POST report\n-> lấy userId + body\n-> gọi ReportService.createReport\n-> trả report",
                "pseudo": "FUNCTION createReport(req,res,next)\n  TRY\n    report = ReportService.createReport(req.user.id,req.body)\n    res.status(201).json(report)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "AdminController.getReports/updateReport",
                "file": "controllers/admin.controller.js",
                "purpose": "Admin xem/cập nhật report.",
                "flow": "GET/PATCH /reports\n-> check MANAGE_REPORTS\n-> gọi ReportService\n-> trả danh sách/report mới",
                "pseudo": "FUNCTION adminReportController(req,res,next)\n  TRY\n    IF GET THEN result = ReportService.getReports(page,limit,filters)\n    IF PATCH THEN result = ReportService.updateReport(id,body,adminId)\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
            {
                "name": "Notification controllers",
                "file": "controllers/admin.controller.js; controllers/user.controller.js",
                "purpose": "Admin gửi, user đọc thông báo.",
                "flow": "Admin POST announcement/reminders\nUser GET/PUT notifications\n-> gọi service tương ứng\n-> trả success/list",
                "pseudo": "FUNCTION notificationController(req,res,next)\n  TRY\n    IF adminSend THEN sendAnnouncement(body)\n    IF dailyReminder THEN createDailyReminders()\n    IF userGet THEN getUserNotifications(userId)\n    IF markRead THEN markNotificationRead(userId,id)\n    res.json(result)\n  CATCH error\n    next(error)\n  END TRY\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "ReportService.createReport/getReports/updateReport",
                "file": "services/report.service.js",
                "purpose": "Lưu và xử lý báo cáo.",
                "flow": "create: validate/infer entity -> insert\nget: filter/paginate\nupdate: đổi status/priority/note",
                "pseudo": "FUNCTION reportService(action)\n  ensureSchema()\n  IF create THEN INSERT Report(userId,entityType,content,status)\n  IF get THEN SELECT Reports WHERE filters OFFSET/FETCH\n  IF update THEN UPDATE Reports SET status,note,adminId\n  RETURN report/result\nEND FUNCTION",
            },
            {
                "name": "AdminService.sendAnnouncement/createDailyReminders",
                "file": "services/admin.service.js",
                "purpose": "Tạo thông báo hệ thống.",
                "flow": "send: xác định audience -> insert notifications\nreminder: tìm user có từ due -> insert reminder\n-> audit log",
                "pseudo": "FUNCTION notificationService(action)\n  IF sendAnnouncement THEN\n    users = query audience\n    FOR user IN users INSERT Notifications\n  IF createDailyReminders THEN\n    learners = query active learners with due words\n    FOR learner INSERT reminder notification\n  logAdminAction\n  RETURN count\nEND FUNCTION",
            },
            {
                "name": "UserService notification methods",
                "file": "services/user.service.js",
                "purpose": "User đọc/đánh dấu thông báo.",
                "flow": "Query notification theo userId\n-> update ReadAt cho một hoặc tất cả\n-> return list/success",
                "pseudo": "FUNCTION userNotificationService(action,userId)\n  IF get THEN SELECT Notifications WHERE UserID=userId\n  IF markRead THEN UPDATE Notifications SET ReadAt=now WHERE id,userId\n  IF markAll THEN UPDATE all unread WHERE userId\n  RETURN result\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module System/Common - Hạ tầng backend",
        "purpose": "Kết nối CSDL, validate request, xử lý lỗi, health-check, integration-test và shutdown an toàn.",
        "components": "config/db.js; validate.js; errorHandler.js; health-check.js; integration-test.js; index.js.",
        "module_flow": "Express app -> middleware JSON/CORS -> routes -> middleware auth/validate -> controller -> service -> DB. Lỗi đi qua next(error) về errorHandler.",
        "controllers": [
            {
                "name": "Express route/middleware chain",
                "file": "index.js; routes/*.js",
                "purpose": "Định tuyến request vào module.",
                "flow": "Request vào Express\n-> CORS/JSON parser\n-> mount route theo prefix\n-> middleware auth/validate\n-> controller\n-> errorHandler nếu lỗi",
                "pseudo": "FUNCTION requestPipeline(req,res)\n  app.use(cors,json)\n  app.use('/api/auth',authRoutes)\n  app.use('/api/user',userRoutes)\n  app.use('/api/admin',adminRoutes)\n  IF error THEN errorHandler(err,req,res,next)\nEND FUNCTION",
            },
        ],
        "services": [
            {
                "name": "DB poolPromise",
                "file": "config/db.js",
                "purpose": "Chia sẻ kết nối SQL Server.",
                "flow": "Đọc biến môi trường\n-> tạo config mssql/msnodesqlv8\n-> tạo poolPromise\n-> export sql,poolPromise",
                "pseudo": "FUNCTION createDbPool()\n  config = read env DB_SERVER, DB_NAME, DB_AUTH\n  poolPromise = sql.connect(config)\n  EXPORT {sql,poolPromise}\nEND FUNCTION",
            },
            {
                "name": "validate(schema)",
                "file": "middlewares/validate.js",
                "purpose": "Validate dữ liệu đầu vào.",
                "flow": "Middleware nhận schema\n-> validate req.body\n-> lỗi trả 400\n-> hợp lệ gán value và next()",
                "pseudo": "FUNCTION validate(schema)\n  RETURN middleware(req,res,next)\n    result = schema.validate(req.body)\n    IF result.error THEN return 400\n    req.body = result.value\n    next()\n  END RETURN\nEND FUNCTION",
            },
            {
                "name": "errorHandler",
                "file": "middlewares/errorHandler.js",
                "purpose": "Chuẩn hóa lỗi API.",
                "flow": "Controller/service throw error\n-> next(error)\n-> errorHandler lấy status/message\n-> trả JSON lỗi",
                "pseudo": "FUNCTION errorHandler(err,req,res,next)\n  status = err.statusCode OR 500\n  message = err.message OR 'Internal Server Error'\n  res.status(status).json({message})\nEND FUNCTION",
            },
            {
                "name": "health-check/integration-test/shutdown",
                "file": "config/health-check.js; config/integration-test.js; index.js",
                "purpose": "Kiểm tra và vận hành server.",
                "flow": "Health: kiểm tra env + DB\nIntegration: query schema/dữ liệu\nShutdown: nhận signal -> close server",
                "pseudo": "FUNCTION opsHelpers()\n  runHealthCheck: check env, connect DB, query sys.tables\n  runIntegrationTest: chạy các query kiểm tra tích hợp\n  gracefulShutdown: server.close rồi process.exit\nEND FUNCTION",
            },
        ],
    },
]


def remove_old_module_section(doc):
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    end = None
    seen_caption = False
    for i, child in enumerate(children):
        if child.tag.split("}")[-1] == "p":
            text = "".join(child.itertext())
            if "Hình 8:" in text:
                seen_caption = True
                continue
            if seen_caption and start is None:
                start = i
            if seen_caption and "THIẾT KẾ DỮ LIỆU" in text:
                end = i
                break
    if start is None or end is None:
        raise RuntimeError("Không tìm thấy vùng module cần thay")
    ref = children[end]
    for child in children[start:end]:
        body.remove(child)
    return ref


def main():
    doc = Document(SRC)
    ref = remove_old_module_section(doc)

    intro = move_paragraph_before(
        doc,
        ref,
        "Các module backend được tách theo lớp Controller và Service. Controller chịu trách nhiệm nhận request, đọc params/body/query và trả response; Service chịu trách nhiệm xử lý nghiệp vụ, truy vấn CSDL, transaction, cập nhật trạng thái và gọi các service liên quan. Mỗi hàm dưới đây đều có luồng xử lý và giả mã.",
    )
    intro.paragraph_format.space_after = Pt(8)

    for idx, module in enumerate(MODULES, 1):
        move_paragraph_before(doc, ref, f"{idx}. {module['name']}", style="Heading 3")
        summary = build_summary_table(doc, module)
        move_table_before(doc, ref, summary)
        move_paragraph_before(doc, ref, "Controller", bold=True)
        controller_table = build_function_table(doc, "Hàm Controller", module["controllers"])
        move_table_before(doc, ref, controller_table)
        move_paragraph_before(doc, ref, "Service / Middleware / Config", bold=True)
        service_table = build_function_table(doc, "Hàm Service", module["services"])
        move_table_before(doc, ref, service_table)
        spacer = move_paragraph_before(doc, ref, "")
        spacer.paragraph_format.space_after = Pt(8)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
