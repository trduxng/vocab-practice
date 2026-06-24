from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("artifacts/design-doc/phuc-report-edit/Detail Design_Phuc_source.docx")
OUT = Path("artifacts/design-doc/phuc-report-edit/Detail Design_Phuc_backend_function_tables_updated.docx")


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, size=9, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05


def add_plain(cell, text, size=9, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_label_para(cell, label, text, size=8.5):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r1 = p.add_run(label)
    set_run_font(r1, size=size, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=size)
    return p


def add_code(cell, code, size=7.8):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code.strip())
    set_run_font(run, size=size, name="Courier New")
    return p


def format_signature(signature):
    if len(signature) <= 34 or "(" not in signature:
        return signature
    name, params = signature.split("(", 1)
    params = params[:-1] if params.endswith(")") else params
    if not params:
        return f"{name}()"
    split_params = params.replace(", ", ",\n  ")
    return f"{name}(\n  {split_params}\n)"


def fill_header(row):
    labels = ["Loại", "Tên hàm / module", "Mô tả chi tiết và giả mã"]
    for idx, cell in enumerate(row.cells):
        clear_cell(cell)
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_plain(cell, labels[idx], size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def set_table_borders(table, color="000000", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def build_clean_table(doc, title, owner, summary, functions):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = [Inches(0.75), Inches(2.35), Inches(3.58)]
    fill_header(table.rows[0])
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)

    row = table.add_row()
    vals = ["Nhóm", title, summary]
    for i, cell in enumerate(row.cells):
        clear_cell(cell)
        shade_cell(cell, "F2F6FA")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        add_plain(cell, vals[i], size=9, bold=(i != 2), align=WD_ALIGN_PARAGRAPH.LEFT)

    row = table.add_row()
    vals = ["Vai trò", owner, "Các hàm bên dưới được chọn vì trực tiếp thể hiện nghiệp vụ cốt lõi, quyền hạn người dùng hoặc luồng xử lý chính của hệ thống VocaBoost. Ký hiệu: S = Service, M = Middleware."]
    for i, cell in enumerate(row.cells):
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        add_plain(cell, vals[i], size=8.8, bold=(i == 0))

    row = table.add_row()
    vals = ["Loại", "Signature", "Chi tiết thực hiện"]
    for i, cell in enumerate(row.cells):
        clear_cell(cell)
        shade_cell(cell, "EAF2F8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        add_plain(cell, vals[i], size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for fn in functions:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            clear_cell(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=110, start=100, bottom=110, end=100)
        kind_label = {"Service": "S", "Middleware": "M", "Controller": "C"}.get(fn["kind"], fn["kind"])
        add_plain(row.cells[0], kind_label, size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_plain(row.cells[1], format_signature(fn["signature"]), size=8.5, bold=True)
        if fn.get("file"):
            add_plain(row.cells[1], fn["file"], size=7.6, italic=True)
        add_label_para(row.cells[2], "Vai trò: ", fn["role"], size=8.4)
        add_plain(row.cells[2], "Giả mã:", size=8.4, bold=True)
        add_code(row.cells[2], fn["pseudo"], size=7.6)
        if fn.get("note"):
            add_label_para(row.cells[2], "Ghi chú: ", fn["note"], size=8.2)
    return table


TABLES = [
    {
        "index": 5,
        "title": "Nhóm xác thực và phân quyền",
        "owner": "Guest, Learner, Creator, Admin",
        "summary": "Nhóm hàm bảo vệ toàn bộ REST API: đăng ký, đăng nhập, sinh JWT, kiểm tra quyền và kiểm tra dữ liệu đầu vào.",
        "functions": [
            {
                "kind": "Service",
                "signature": "AuthService.register(fullName, email, password)",
                "file": "backend/src/services/auth.service.js",
                "role": "Tạo tài khoản mới cho người học và gán vai trò mặc định Learner.",
                "pseudo": """
FUNCTION register(fullName, email, password)
    Kết nối CSDL
    Tìm user có Email = email
    IF email đã tồn tại THEN
        Báo lỗi "Email already registered"
    END IF
    Hash password bằng bcrypt
    Lấy RoleID của vai trò Learner
    INSERT Users(fullName, email, passwordHash, roleId, isActive)
    Trả về thông tin user mới
END FUNCTION
""",
                "note": "Hàm này thuộc luồng Guest -> Learner, không trả passwordHash ra client.",
            },
            {
                "kind": "Service",
                "signature": "AuthService.login(email, password)",
                "file": "backend/src/services/auth.service.js",
                "role": "Xác thực đăng nhập và cấp JWT token để frontend gọi các API được bảo vệ.",
                "pseudo": """
FUNCTION login(email, password)
    Tìm user đang active theo email
    IF không tìm thấy THEN báo lỗi 401
    So sánh password với PasswordHash bằng bcrypt
    IF mật khẩu sai THEN báo lỗi 401
    Lấy danh sách permission theo RoleID
    Tạo JWT chứa id, email, role, permissions
    IF user là Learner THEN
        Thưởng daily login XP nếu đủ điều kiện
    END IF
    Trả về token và thông tin user
END FUNCTION
""",
                "note": "Token là cơ sở cho middleware verifyToken và phân quyền động.",
            },
            {
                "kind": "Middleware",
                "signature": "verifyToken(req, res, next)",
                "file": "backend/src/middlewares/auth.js",
                "role": "Kiểm tra request có JWT hợp lệ trước khi cho truy cập API nội bộ.",
                "pseudo": """
FUNCTION verifyToken(req, res, next)
    Đọc Authorization header
    IF header không có Bearer token THEN
        Trả về HTTP 401
    END IF
    Giải mã token bằng JWT_SECRET
    IF token sai hoặc hết hạn THEN
        Trả về HTTP 401
    END IF
    Gán thông tin user vào req.user
    Chuyển sang middleware/controller tiếp theo
END FUNCTION
""",
            },
            {
                "kind": "Middleware",
                "signature": "checkPermission(permission)",
                "file": "backend/src/middlewares/auth.js",
                "role": "Bảo vệ endpoint chỉ cho user có đúng quyền cụ thể, ví dụ MANAGE_WORDS hoặc VIEW_DASHBOARD.",
                "pseudo": """
FUNCTION checkPermission(permission)
    RETURN middleware(req, res, next)
        permissions = req.user.permissions
        IF permissions không chứa permission THEN
            Trả về HTTP 403
        END IF
        Cho phép đi tiếp
    END RETURN
END FUNCTION
""",
            },
            {
                "kind": "Middleware",
                "signature": "checkAnyPermission(permissions)",
                "file": "backend/src/middlewares/auth.js",
                "role": "Cho phép một endpoint được truy cập bởi nhiều nhóm quyền khác nhau.",
                "pseudo": """
FUNCTION checkAnyPermission(permissions)
    RETURN middleware(req, res, next)
        userPermissions = req.user.permissions
        hasPermission = tồn tại quyền trong permissions mà user có
        IF hasPermission = false THEN
            Trả về HTTP 403
        END IF
        Cho phép đi tiếp
    END RETURN
END FUNCTION
""",
            },
            {
                "kind": "Middleware",
                "signature": "validate(schema)",
                "file": "backend/src/middlewares/validate.js",
                "role": "Kiểm tra dữ liệu request body/query/params trước khi vào controller.",
                "pseudo": """
FUNCTION validate(schema)
    RETURN middleware(req, res, next)
        Chạy schema.validate(req.body)
        IF dữ liệu không hợp lệ THEN
            Trả về HTTP 400 kèm danh sách lỗi
        END IF
        Gán dữ liệu đã chuẩn hóa vào req.body
        Cho phép đi tiếp
    END RETURN
END FUNCTION
""",
            },
        ],
    },
    {
        "index": 6,
        "title": "Nhóm người học: chủ đề, lộ trình, SRS, mini-test và tiến độ",
        "owner": "Learner",
        "summary": "Nhóm hàm quan trọng nhất của dự án, thể hiện quá trình học từ vựng TOEIC, ôn tập lặp lại ngắt quãng và theo dõi tiến độ cá nhân.",
        "functions": [
            {
                "kind": "Service",
                "signature": "CategoriesService.getTopics(userId)",
                "file": "backend/src/services/categories.service.js",
                "role": "Hiển thị danh sách chủ đề kèm thống kê tiến độ học của từng người dùng.",
                "pseudo": """
FUNCTION getTopics(userId)
    Kết nối CSDL
    Lấy danh sách Topics đang Published
    JOIN WordTopics và Words để đếm wordCount
    IF có userId THEN
        JOIN UserWordProgress theo userId
        Tính learnedCount, masteredCount, dueCount
        Tính averageMastery và progressPercent
    ELSE
        Trả về thống kê cơ bản của topic
    END IF
    Sắp xếp theo displayOrder/topicName
    Trả về danh sách topics
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "LearningPathService.getRoadmap(userId)",
                "file": "backend/src/services/learning-path.service.js",
                "role": "Tạo lộ trình học theo các level TOEIC và trạng thái hoàn thành của learner.",
                "pseudo": """
FUNCTION getRoadmap(userId)
    Đảm bảo schema LearningPath đã tồn tại
    Đồng bộ các topic Published vào roadmap
    Lấy danh sách level và topic thuộc từng level
    Lấy tiến độ học của user nếu có userId
    Gọi buildRoadmap để gom nhóm dữ liệu
    Trả về roadmap cho frontend hiển thị
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "UserService.getDueFlashcards(userId, options)",
                "file": "backend/src/services/user.service.js",
                "role": "Lấy danh sách flashcard cần học/ôn dựa trên lịch SRS và bộ lọc topic.",
                "pseudo": """
FUNCTION getDueFlashcards(userId, options)
    Đọc topicId và mode từ options
    Lấy giới hạn ôn tập của user
    Query Words Published
    LEFT JOIN UserWordProgress theo userId
    IF mode = "new" THEN lấy từ chưa học
    IF mode = "daily" THEN lấy từ có NextReviewDate <= hiện tại
    Ưu tiên từ đến hạn, từ yếu, sau đó từ mới
    Trả về danh sách flashcards
END FUNCTION
""",
                "note": "Đây là nơi thể hiện rõ kỹ thuật Spaced Repetition ở bước chọn từ cần ôn.",
            },
            {
                "kind": "Service",
                "signature": "UserService.submitAnswer(payload)",
                "file": "backend/src/services/user.service.js",
                "role": "Ghi nhận đáp án câu hỏi, cập nhật phản hồi ôn tập và cộng XP nếu là hoạt động học từ.",
                "pseudo": """
FUNCTION submitAnswer(payload)
    Nhận userId, questionId, submittedAnswer, reviewRating
    Gọi stored procedure usp_SubmitQuestionAttempt
    Lấy WordID chuẩn từ kết quả câu hỏi
    IF không lấy được WordID THEN báo lỗi
    IF reviewRating thuộc Again/Hard/Good/Easy THEN
        Cập nhật EaseFactor và NextReviewDate
        Again -> ôn lại sau 10 phút
        Hard -> ôn lại sau 1 ngày
        Good/Easy -> giãn ngày theo MasteryLevel
    END IF
    IF activityType = LearnWord THEN cộng XP
    Trả về kết quả chấm, lịch ôn tiếp theo và XP
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "UserService.submitWordReview(payload)",
                "file": "backend/src/services/user.service.js",
                "role": "Cập nhật trạng thái ghi nhớ của một từ sau khi learner tự đánh giá flashcard.",
                "pseudo": """
FUNCTION submitWordReview(payload)
    Nhận userId, wordId, isCorrect, reviewRating
    MERGE UserWordProgress theo userId + wordId
    IF đã có tiến độ THEN
        Tăng/giảm MasteryLevel theo isCorrect
        Cập nhật ConsecutiveCorrect/ConsecutiveWrong
        Tính NextReviewDate theo reviewRating
        Điều chỉnh EaseFactor trong khoảng cho phép
        Cập nhật MemoryStatus: Learning/Reviewing/Mastered/Lapsed
    ELSE
        Tạo bản ghi tiến độ ban đầu cho từ
    END IF
    IF activityType = LearnWord THEN cộng XP
    Trả về masteryLevel, memoryStatus, nextReviewDate
END FUNCTION
""",
                "note": "Hàm này là lõi của SRS engine trong dự án.",
            },
            {
                "kind": "Service",
                "signature": "UserService.getSmartReviewQueue(userId, limit)",
                "file": "backend/src/services/user.service.js",
                "role": "Tạo hàng đợi ôn tập thông minh, ưu tiên từ sắp quên hoặc sai nhiều.",
                "pseudo": """
FUNCTION getSmartReviewQueue(userId, limit)
    Chuẩn hóa limit trong khoảng 1..50
    Lấy Words Published đã có UserWordProgress
    Chỉ lấy từ có NextReviewDate trong 7 ngày tới
    Tính priorityScore:
        quá hạn nhiều giờ -> điểm ưu tiên cao
        có ConsecutiveWrong -> nhân hệ số ưu tiên
        chưa đến hạn -> ưu tiên thấp hơn
    Sắp xếp theo priorityScore và masteryLevel
    Trả về danh sách từ cần ôn
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "UserService.submitMiniTestBatch(userId, testId, answers)",
                "file": "backend/src/services/user.service.js",
                "role": "Chấm toàn bộ mini-test, lưu lịch sử làm bài và cập nhật tiến độ từng từ.",
                "pseudo": """
FUNCTION submitMiniTestBatch(userId, testId, answers)
    Bắt đầu transaction
    Lấy danh sách câu hỏi thuộc miniTest
    IF miniTest không có câu hỏi THEN báo lỗi
    IF số đáp án không khớp THEN báo lỗi
    FOR mỗi answer
        Kiểm tra question có thuộc miniTest không
        So sánh submittedAnswer với CorrectAnswer
        INSERT ExerciseAttempts
        MERGE UserWordProgress cho WordID tương ứng
        Cộng correctCount nếu đúng
    END FOR
    INSERT MiniTestAttempts với điểm tổng
    Commit transaction
    Cộng XP hoàn thành mini-test
    Trả về score, correctCount và chi tiết từng câu
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "UserService.getProgressAnalytics(userId)",
                "file": "backend/src/services/user.service.js",
                "role": "Tổng hợp dữ liệu biểu đồ dashboard học tập của người học.",
                "pseudo": """
FUNCTION getProgressAnalytics(userId)
    Lấy tổng số từ đã học và đã mastered
    Tính phân bố MasteryLevel
    Tính activity theo ngày/tuần/tháng
    Tính tiến độ theo topic/level
    Lấy accuracy từ ExerciseAttempts
    Gom dữ liệu thành analytics object
    Trả về dữ liệu cho Learning Dashboard
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "UserService.getMistakeReviewQueue(userId, limit)",
                "file": "backend/src/services/user.service.js",
                "role": "Lấy danh sách từ/câu hỏi learner hay sai để ôn lại có trọng tâm.",
                "pseudo": """
FUNCTION getMistakeReviewQueue(userId, limit)
    Chuẩn hóa limit
    Query ExerciseAttempts của user
    Lọc các attempt sai
    GROUP BY WordID/QuestionID để đếm số lần sai
    JOIN Words và Questions để lấy nội dung ôn
    Sắp xếp theo số lỗi và thời điểm sai gần nhất
    Trả về danh sách lỗi cần ôn
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "UserService.addNotebookEntry(userId, data)",
                "file": "backend/src/services/user.service.js",
                "role": "Cho phép learner lưu từ vào sổ tay cá nhân kèm ghi chú riêng.",
                "pseudo": """
FUNCTION addNotebookEntry(userId, data)
    Nhận wordId, note, isFavorite
    Kiểm tra word tồn tại
    Kiểm tra entry đã có trong notebook chưa
    IF đã tồn tại THEN cập nhật ghi chú/isFavorite
    ELSE INSERT UserVocabularyNotebook
    END IF
    Trả về entry mới nhất
END FUNCTION
""",
            },
        ],
    },
    {
        "index": 7,
        "title": "Nhóm quản trị và biên tập nội dung",
        "owner": "Admin, Content Creator",
        "summary": "Nhóm hàm phục vụ quản lý kho từ vựng TOEIC, câu hỏi, mini-test, import dữ liệu và phân quyền người dùng.",
        "functions": [
            {
                "kind": "Service",
                "signature": "AdminService.getDashboardStats()",
                "file": "backend/src/services/admin.service.js",
                "role": "Tổng hợp số liệu cho trang Admin Dashboard.",
                "pseudo": """
FUNCTION getDashboardStats()
    Đếm tổng Users, Words, Questions, Topics, MiniTests
    Đếm số nội dung PendingReview
    Đếm số learner đang hoạt động
    Tính accuracy/attempt trung bình nếu có dữ liệu
    Gom các chỉ số thành object thống kê
    Trả về dashboard stats
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.getTopics(page, limit, filters)",
                "file": "backend/src/services/admin.service.js",
                "role": "Quản trị danh sách chủ đề, hỗ trợ tìm kiếm, lọc trạng thái và phân trang.",
                "pseudo": """
FUNCTION getTopics(page, limit, filters)
    Chuẩn hóa page và limit
    Tạo điều kiện WHERE theo search/status/categoryId
    Query Topics kèm CategoryName
    Đếm wordCount và miniTestCount của từng topic
    Áp dụng ORDER BY và OFFSET/FETCH
    Query tổng số bản ghi
    Trả về { data, total, page, limit }
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.createWord(wordData, adminId)",
                "file": "backend/src/services/admin.service.js",
                "role": "Admin thêm từ vựng TOEIC mới cùng topic và câu ví dụ.",
                "pseudo": """
FUNCTION createWord(wordData, adminId)
    Kiểm tra ContentStatus hợp lệ
    Chuẩn hóa topicIds và examples
    Bắt đầu transaction
    INSERT Words(term, meaning, phonetic, partOfSpeechId, status)
    FOR mỗi topicId
        INSERT WordTopics(wordId, topicId)
    END FOR
    FOR mỗi example hợp lệ
        INSERT ExampleSentences(wordId, sentence, translation)
    END FOR
    Commit transaction
    Ghi AdminAuditLog CREATE_WORD
    Trả về id và thông tin từ vừa tạo
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.updateWord(wordId, wordData, adminId)",
                "file": "backend/src/services/admin.service.js",
                "role": "Cập nhật từ vựng, topic liên kết và ví dụ minh họa.",
                "pseudo": """
FUNCTION updateWord(wordId, wordData, adminId)
    Kiểm tra status hợp lệ
    Chuẩn hóa topicIds/examples nếu được gửi lên
    Bắt đầu transaction
    UPDATE Words theo wordId
    IF không có dòng bị cập nhật THEN rollback và trả false
    IF có topicIds THEN
        Xóa WordTopics cũ
        Thêm lại WordTopics mới
    END IF
    IF có examples THEN
        Xóa ExampleSentences cũ
        Thêm lại examples mới
    END IF
    Commit transaction
    Ghi audit log UPDATE_WORD
    Trả về true
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.previewWordImport(input)",
                "file": "backend/src/services/admin.service.js",
                "role": "Đọc trước file CSV/plain text để admin kiểm tra lỗi trước khi import thật.",
                "pseudo": """
FUNCTION previewWordImport(input)
    Parse nội dung CSV/text thành các dòng dữ liệu
    FOR mỗi dòng
        Chuẩn hóa key: term, meaning, phonetic, topics, examples
        Kiểm tra trường bắt buộc
        Kiểm tra trùng từ hoặc dữ liệu sai định dạng
        Đưa vào danh sách validRows hoặc errors
    END FOR
    Trả về preview gồm tổng dòng, dòng hợp lệ, dòng lỗi
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.bulkInsertWords(input, adminId)",
                "file": "backend/src/services/admin.service.js",
                "role": "Import hàng loạt từ vựng để xây dựng nhanh ngân hàng TOEIC.",
                "pseudo": """
FUNCTION bulkInsertWords(input, adminId)
    Parse dữ liệu import
    Bắt đầu transaction
    FOR mỗi word hợp lệ
        INSERT Words
        Tạo Topic nếu cần hoặc lấy topic hiện có
        INSERT WordTopics
        INSERT ExampleSentences nếu có
    END FOR
    Nếu có lỗi nghiêm trọng THEN rollback
    Commit transaction
    Ghi audit log BULK_IMPORT_WORDS
    Trả về số dòng thành công/thất bại
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.createMiniTest(testData, adminId)",
                "file": "backend/src/services/admin.service.js",
                "role": "Admin tạo bài kiểm tra nhanh theo topic và danh sách câu hỏi.",
                "pseudo": """
FUNCTION createMiniTest(testData, adminId)
    Kiểm tra topic và danh sách questionIds
    Bắt đầu transaction
    INSERT MiniTests(title, description, topicId, status)
    FOR mỗi questionId
        INSERT MiniTestItems(miniTestId, questionId, displayOrder)
    END FOR
    Commit transaction
    Ghi audit log CREATE_MINITEST
    Trả về miniTest vừa tạo
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.updateUserRole(userId, roleName)",
                "file": "backend/src/services/admin.service.js",
                "role": "Admin đổi vai trò người dùng giữa Learner, Creator hoặc Admin.",
                "pseudo": """
FUNCTION updateUserRole(userId, roleName)
    Kiểm tra roleName hợp lệ
    Tìm RoleID theo roleName
    IF role không tồn tại THEN báo lỗi
    UPDATE Users SET RoleID = roleId
    IF không có user bị cập nhật THEN báo lỗi
    Trả về user với role mới
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "CreatorService.createWord(data, userId)",
                "file": "backend/src/services/creator.service.js",
                "role": "Content Creator tạo từ vựng ở trạng thái Draft để gửi duyệt.",
                "pseudo": """
FUNCTION createWord(data, userId)
    Chuẩn hóa topicIds
    Bắt đầu transaction
    Kiểm tra PartOfSpeech tồn tại
    Kiểm tra từ chưa bị trùng theo Term + PartOfSpeech
    FOR mỗi topicId
        Kiểm tra creator được dùng topic đó
    END FOR
    INSERT Words với ContentStatus = Draft
    INSERT WordTopics và ExampleSentences
    Commit transaction
    Trả về id từ vừa tạo
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "CreatorService.createQuestion(data, userId)",
                "file": "backend/src/services/creator.service.js",
                "role": "Content Creator tạo câu hỏi luyện tập gắn với từ vựng.",
                "pseudo": """
FUNCTION createQuestion(data, userId)
    Kiểm tra WordID tồn tại và creator được sử dụng
    Kiểm tra dạng câu hỏi, đáp án đúng và options
    Bắt đầu transaction
    INSERT Questions với ContentStatus = Draft
    Nếu có options thì INSERT QuestionOptions
    Commit transaction
    Trả về question vừa tạo
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "CreatorService.createMiniTest(data, userId)",
                "file": "backend/src/services/creator.service.js",
                "role": "Content Creator tạo mini-test trước khi gửi admin xét duyệt.",
                "pseudo": """
FUNCTION createMiniTest(data, userId)
    Kiểm tra topic và questionIds hợp lệ
    Bắt đầu transaction
    INSERT MiniTests với ContentStatus = Draft
    FOR mỗi questionId
        INSERT MiniTestItems
    END FOR
    Commit transaction
    Trả về mini-test mới
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "CreatorService.submitForReview(tableName, idColumn, id, userId)",
                "file": "backend/src/services/creator.service.js",
                "role": "Chuyển nội dung của creator từ Draft/Rejected sang PendingReview.",
                "pseudo": """
FUNCTION submitForReview(tableName, idColumn, id, userId)
    Kiểm tra entity thuộc user hiện tại
    IF ContentStatus không phải Draft hoặc Rejected THEN báo lỗi
    UPDATE entity SET ContentStatus = PendingReview
    Ghi thời điểm cập nhật
    Trả về nội dung đã gửi duyệt
END FUNCTION
""",
            },
        ],
    },
    {
        "index": 8,
        "title": "Nhóm duyệt nội dung, AI, gamification và thông báo",
        "owner": "Admin, Creator, Learner",
        "summary": "Nhóm hàm hỗ trợ vận hành hệ thống: duyệt nội dung, sinh gợi ý AI, cộng điểm XP, mở khóa thành tích và gửi thông báo.",
        "functions": [
            {
                "kind": "Service",
                "signature": "ReviewService.getPendingContent()",
                "file": "backend/src/services/review.service.js",
                "role": "Admin lấy toàn bộ nội dung đang chờ duyệt từ nhiều bảng.",
                "pseudo": """
FUNCTION getPendingContent()
    Query Topics có ContentStatus = PendingReview
    UNION với Words PendingReview
    UNION với Questions PendingReview
    UNION với MiniTests PendingReview
    JOIN Users để lấy creatorName
    Sắp xếp theo createdAt tăng dần
    Trả về danh sách nội dung chờ duyệt
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "ReviewService.approve(entityType, entityId, adminId)",
                "file": "backend/src/services/review.service.js",
                "role": "Admin phê duyệt nội dung, chuyển từ PendingReview sang Published.",
                "pseudo": """
FUNCTION approve(entityType, entityId, adminId)
    Resolve entityType thành table và khóa chính
    Bắt đầu transaction
    UPDATE entity SET ContentStatus = Published
    Cập nhật ReviewedByUserID, ReviewedAt, PublishedAt nếu có
    IF entity là MiniTest THEN set IsPublished = 1
    INSERT ContentReviewLogs với NewStatus = Published
    Commit transaction
    Trả về true nếu duyệt thành công
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "ReviewService.reject(entityType, entityId, adminId, reason)",
                "file": "backend/src/services/review.service.js",
                "role": "Admin từ chối nội dung và lưu lý do để creator chỉnh sửa.",
                "pseudo": """
FUNCTION reject(entityType, entityId, adminId, reason)
    Resolve entityType thành table và khóa chính
    Bắt đầu transaction
    UPDATE entity SET ContentStatus = Rejected
    Cập nhật ReviewedByUserID và ReviewedAt
    INSERT ContentReviewLogs với comment = reason
    Commit transaction
    Trả về true nếu từ chối thành công
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AiService.suggestWordContent(payload)",
                "file": "backend/src/services/ai.service.js",
                "role": "Gợi ý nội dung từ vựng TOEIC cho creator/admin: nghĩa, ví dụ, giải thích.",
                "pseudo": """
FUNCTION suggestWordContent(payload)
    Nhận term, meaning, partOfSpeech, exampleCount
    Làm sạch dữ liệu đầu vào
    Tạo prompt yêu cầu sinh nội dung theo ngữ cảnh TOEIC
    Gọi AI provider với timeout
    Trích xuất response text
    Parse JSON từ response
    Chuẩn hóa danh sách ví dụ
    Trả về nội dung gợi ý cho frontend
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "GamificationService.awardXP(userId, event)",
                "file": "backend/src/services/gamification.service.js",
                "role": "Cộng điểm kinh nghiệm cho learner sau các hoạt động học tập.",
                "pseudo": """
FUNCTION awardXP(userId, event)
    Đảm bảo schema gamification tồn tại
    Xác định số XP theo eventType
    IF sourceKey đã được thưởng trước đó THEN
        Không cộng trùng XP
    END IF
    INSERT UserXPEvents
    Cập nhật TotalXP của Users
    Tính level mới từ tổng XP
    Kiểm tra achievement mới
    Trả về xpGained, totalXP, levelState
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "GamificationService.checkAchievements(userId)",
                "file": "backend/src/services/gamification.service.js",
                "role": "Kiểm tra điều kiện mở khóa huy hiệu/thành tích của learner.",
                "pseudo": """
FUNCTION checkAchievements(userId)
    Lấy metrics học tập của user
    Lấy danh sách achievements hiện có
    FOR mỗi achievement
        Kiểm tra điều kiện: XP, streak, số từ mastered, mini-test
        IF đạt điều kiện và chưa nhận THEN
            INSERT UserAchievements
            Đưa vào danh sách newlyUnlocked
        END IF
    END FOR
    Trả về achievement mới mở khóa
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.sendAnnouncement(data, adminId)",
                "file": "backend/src/services/admin.service.js",
                "role": "Admin gửi thông báo hệ thống cho learner/creator hoặc toàn bộ người dùng.",
                "pseudo": """
FUNCTION sendAnnouncement(data, adminId)
    Kiểm tra title, message, targetRole/targetUser
    Xác định danh sách người nhận
    FOR mỗi user nhận thông báo
        INSERT Notifications(userId, title, message, type)
    END FOR
    Ghi audit log SEND_ANNOUNCEMENT
    Trả về số thông báo đã tạo
END FUNCTION
""",
            },
            {
                "kind": "Service",
                "signature": "AdminService.createDailyReminders(adminId)",
                "file": "backend/src/services/admin.service.js",
                "role": "Tạo nhắc nhở học tập hằng ngày cho người học còn từ đến hạn ôn.",
                "pseudo": """
FUNCTION createDailyReminders(adminId)
    Tìm learner đang active
    Với mỗi learner, đếm từ có NextReviewDate <= hiện tại
    IF learner có từ cần ôn THEN
        INSERT Notification nhắc học hôm nay
    END IF
    Ghi audit log CREATE_DAILY_REMINDERS
    Trả về số reminder đã tạo
END FUNCTION
""",
            },
        ],
    },
]


def main():
    doc = Document(SRC)
    targets = {spec["index"]: doc.tables[spec["index"]] for spec in TABLES}
    for spec in TABLES:
        old_table = targets[spec["index"]]
        new_table = build_clean_table(
            doc,
            spec["title"],
            spec["owner"],
            spec["summary"],
            spec["functions"],
        )
        old_table._tbl.addnext(new_table._tbl)
        old_table._tbl.getparent().remove(old_table._tbl)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
