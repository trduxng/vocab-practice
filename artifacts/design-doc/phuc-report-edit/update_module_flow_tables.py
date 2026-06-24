from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("artifacts/design-doc/phuc-report-edit/Detail Design_Phuc_backend_function_tables_updated.docx")
OUT = Path("artifacts/design-doc/phuc-report-edit/Detail Design_Phuc_module_flow_updated.docx")


def set_run_font(run, size=9, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05


def set_cell_margins(cell, top=90, start=95, bottom=90, end=95):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="000000", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_text(cell, text, size=8.7, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_label(cell, label, text, size=8.5):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r1 = p.add_run(label)
    set_run_font(r1, size=size, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=size)


def add_code(cell, text, size=7.5):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text.strip())
    set_run_font(r, size=size, name="Courier New")


def insert_paragraph_before(ref, text="", style=None):
    p = OxmlElement("w:p")
    ref.addprevious(p)
    paragraph = None
    # python-docx has no direct wrapper constructor exposed on Document.
    # Add a paragraph at the end, move its element before ref, then fill it.
    return p


def move_paragraph_before(doc, ref, text, style=None):
    paragraph = doc.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    ref.addprevious(paragraph._p)
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        size = 11 if style == "Heading 3" else 9
        set_run_font(run, size=size, bold=style in ("Heading 2", "Heading 3"))
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def move_table_before(doc, ref, table):
    table._tbl.getparent().remove(table._tbl)
    ref.addprevious(table._tbl)


def build_module_table(doc, module):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = [Inches(2.15), Inches(1.75), Inches(2.78)]

    headers = ["Thành phần / hàm", "Vai trò", "Luồng xử lý"]
    for i, cell in enumerate(table.rows[0].cells):
        clear_cell(cell)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(cell, headers[i], size=9.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Module summary
    row = table.add_row()
    for i, cell in enumerate(row.cells):
        clear_cell(cell)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        shade_cell(cell, "F2F6FA")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_text(row.cells[0], module["name"], size=9, bold=True)
    add_label(row.cells[1], "Mục đích: ", module["purpose"], size=8.4)
    add_label(row.cells[2], "Thành phần: ", module["components"], size=8.4)
    add_label(row.cells[2], "Luồng module: ", module["flow"], size=8.4)

    # Function header
    row = table.add_row()
    labels = ["Hàm/API chính", "Chức năng", "Luồng xử lý từng hàm"]
    for i, cell in enumerate(row.cells):
        clear_cell(cell)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        shade_cell(cell, "EAF2F8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(cell, labels[i], size=8.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for fn in module["functions"]:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            clear_cell(cell)
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        add_text(row.cells[0], fn["name"], size=8.3, bold=True)
        if fn.get("file"):
            add_text(row.cells[0], fn["file"], size=7.4, italic=True)
        add_text(row.cells[1], fn["role"], size=8.2)
        add_code(row.cells[2], fn["flow"], size=7.2)

    return table


MODULES = [
    {
        "name": "Module Auth - Xác thực và phân quyền",
        "purpose": "Đăng ký, đăng nhập, phát hành JWT và chuẩn bị thông tin quyền cho các request sau.",
        "components": "auth.routes.js; AuthController; AuthService; validate middleware; auth middleware; bảng Users, Roles, Permissions.",
        "flow": "Client -> /api/auth -> validate(schema) -> AuthController -> AuthService -> SQL Server -> JSON response. Các API nội bộ đi qua verifyToken/checkPermission trước controller.",
        "functions": [
            {
                "name": "POST /register\nAuthController.register\nAuthService.register",
                "file": "controllers/auth.controller.js; services/auth.service.js",
                "role": "Tạo tài khoản Learner mới.",
                "flow": "FUNCTION register\n  Nhận fullName, email, password\n  Validate body theo schemas.register\n  Controller gọi AuthService.register\n  Service kiểm tra email trùng\n  Hash password bằng bcrypt\n  Lấy RoleID Learner\n  INSERT Users\n  Trả HTTP 201 + user\nEND FUNCTION",
            },
            {
                "name": "POST /login\nAuthController.login\nAuthService.login",
                "file": "controllers/auth.controller.js; services/auth.service.js",
                "role": "Xác thực email/password và cấp JWT.",
                "flow": "FUNCTION login\n  Validate email/password\n  Tìm user active theo email\n  So khớp password bằng bcrypt.compare\n  Lấy role và permissions\n  Tạo JWT chứa id, role, permissions\n  Nếu Learner thì awardDailyLogin\n  Trả HTTP 200 + token + user\nEND FUNCTION",
            },
            {
                "name": "verifyToken(req,res,next)",
                "file": "middlewares/auth.js",
                "role": "Bảo vệ API cần đăng nhập.",
                "flow": "FUNCTION verifyToken\n  Đọc Authorization Bearer token\n  Nếu thiếu token -> 401\n  Verify token bằng JWT_SECRET\n  Gán payload vào req.user\n  next()\nEND FUNCTION",
            },
            {
                "name": "checkPermission / checkAnyPermission",
                "file": "middlewares/auth.js",
                "role": "Kiểm tra quyền trước khi vào controller.",
                "flow": "FUNCTION checkPermission\n  Lấy req.user.permissions\n  Nếu thiếu quyền yêu cầu -> 403\n  Nếu hợp lệ -> next()\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Categories - Danh mục học tập công khai",
        "purpose": "Cung cấp loại từ và danh sách chủ đề học cho frontend, có kèm tiến độ nếu user đã đăng nhập.",
        "components": "categories.routes.js; CategoriesController; CategoriesService; bảng PartOfSpeeches, Topics, WordTopics, UserWordProgress.",
        "flow": "Client -> /api/categories -> CategoriesController -> CategoriesService -> SQL Server -> danh sách danh mục/chủ đề.",
        "functions": [
            {
                "name": "GET /part-of-speeches\ngetPartOfSpeeches()",
                "file": "controllers/categories.controller.js; services/categories.service.js",
                "role": "Lấy danh sách loại từ để hiển thị hoặc dùng khi tạo từ vựng.",
                "flow": "FUNCTION getPartOfSpeeches\n  Controller nhận request\n  Gọi CategoriesService.getPartOfSpeeches\n  Query bảng PartOfSpeeches\n  ORDER BY tên loại từ\n  Trả HTTP 200 + recordset\nEND FUNCTION",
            },
            {
                "name": "GET /topics\ngetTopics(userId)",
                "file": "controllers/categories.controller.js; services/categories.service.js",
                "role": "Lấy topic kèm wordCount, learnedCount, masteredCount, dueCount, progressPercent.",
                "flow": "FUNCTION getTopics\n  Lấy userId nếu request đã có token\n  Query Topics Published\n  JOIN WordTopics/Words để đếm số từ\n  LEFT JOIN UserWordProgress theo userId\n  Tính learned/mastered/due/progress\n  Trả danh sách topic cho frontend\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module User Learning - Học từ vựng, SRS, mini-test, notebook",
        "purpose": "Module nghiệp vụ chính cho Learner: học flashcard, gửi đáp án, ôn tập thông minh, làm mini-test, xem thống kê và quản lý sổ tay.",
        "components": "user.routes.js; UserController; UserService; GamificationService; LearningPathController; SQL tables Words, Questions, UserWordProgress, ExerciseAttempts, MiniTestAttempts, Notebook.",
        "flow": "Client Learner -> /api/user -> verifyToken -> UserController -> UserService -> SQL Server. Khi có hoạt động học, service cập nhật UserWordProgress và gọi GamificationService để cộng XP.",
        "functions": [
            {
                "name": "GET /flashcards\ngetFlashcards / getDueFlashcards",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Lấy flashcard mới hoặc đến hạn ôn theo SRS.",
                "flow": "FUNCTION getFlashcards\n  Đọc query topicId/mode\n  Lấy userId từ req.user\n  Gọi getDueFlashcards(userId, options)\n  Query Words Published + UserWordProgress\n  Ưu tiên từ quá hạn, từ yếu, từ mới\n  Trả danh sách flashcards\nEND FUNCTION",
            },
            {
                "name": "POST /submit-answer\nsubmitAnswer(payload)",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Chấm câu hỏi và cập nhật lịch ôn khi người học trả lời.",
                "flow": "FUNCTION submitAnswer\n  Nhận questionId, submittedAnswer, reviewRating\n  Gọi usp_SubmitQuestionAttempt để ghi attempt\n  Lấy WordID của câu hỏi\n  Nếu rating Again/Hard/Good/Easy\n    cập nhật EaseFactor, NextReviewDate\n  Nếu activityType=LearnWord thì awardXP\n  Trả kết quả đúng/sai + lịch ôn mới\nEND FUNCTION",
            },
            {
                "name": "submitWordReview(payload)",
                "file": "services/user.service.js",
                "role": "Cập nhật tiến độ một từ sau khi lật flashcard.",
                "flow": "FUNCTION submitWordReview\n  MERGE UserWordProgress theo userId+wordId\n  Nếu đúng -> tăng MasteryLevel, ConsecutiveCorrect\n  Nếu sai -> giảm MasteryLevel, tăng ConsecutiveWrong\n  Tính NextReviewDate theo Again/Hard/Good/Easy\n  Cập nhật MemoryStatus Learning/Reviewing/Mastered/Lapsed\n  Trả masteryLevel và nextReviewDate\nEND FUNCTION",
            },
            {
                "name": "GET /review/smart-queue\ngetSmartReviewQueue",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Tạo hàng đợi ôn tập ưu tiên từ sắp quên hoặc sai nhiều.",
                "flow": "FUNCTION getSmartReviewQueue\n  Chuẩn hóa limit\n  Query UserWordProgress có NextReviewDate trong 7 ngày\n  Tính priorityScore theo độ quá hạn và ConsecutiveWrong\n  ORDER BY priorityScore DESC, MasteryLevel ASC\n  Trả danh sách từ cần ôn\nEND FUNCTION",
            },
            {
                "name": "GET /topics/:topicId/words\ngetTopicWords",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Hiển thị toàn bộ từ trong một chủ đề và tiến độ từng từ.",
                "flow": "FUNCTION getTopicWords\n  Nhận topicId từ params\n  Query Words JOIN WordTopics\n  LEFT JOIN UserWordProgress theo userId\n  Lấy câu ví dụ liên quan\n  Trả danh sách word + progress\nEND FUNCTION",
            },
            {
                "name": "GET /minitests\nGET /minitests/:id\ngetMiniTests/getMiniTestDetails",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Lấy danh sách mini-test và chi tiết câu hỏi.",
                "flow": "FUNCTION getMiniTests\n  Đọc phân trang\n  Query MiniTests Published\n  Đếm số câu và attempt của user\n  Trả danh sách bài test\nEND FUNCTION\n\nFUNCTION getMiniTestDetails\n  Nhận testId\n  Query MiniTests + MiniTestItems + Questions\n  Trả chi tiết bài test\nEND FUNCTION",
            },
            {
                "name": "POST /minitests/:id/submit\nsubmitMiniTestBatch",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Nộp mini-test, chấm điểm và cập nhật tiến độ từ vựng.",
                "flow": "FUNCTION submitMiniTestBatch\n  Bắt đầu transaction\n  Lấy danh sách câu hỏi của mini-test\n  Kiểm tra số đáp án và questionId hợp lệ\n  FOR mỗi đáp án\n    So sánh với CorrectAnswer\n    INSERT ExerciseAttempts\n    MERGE UserWordProgress\n  END FOR\n  INSERT MiniTestAttempts\n  Commit\n  Cộng XP hoàn thành bài test\n  Trả score và chi tiết kết quả\nEND FUNCTION",
            },
            {
                "name": "GET /progress/analytics\ngetProgressAnalytics",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Tổng hợp dữ liệu dashboard học tập.",
                "flow": "FUNCTION getProgressAnalytics\n  Query tổng từ đã học/đã mastered\n  Tính accuracy từ ExerciseAttempts\n  Tính tiến độ theo topic/level\n  Tính hoạt động theo ngày/tuần/tháng\n  Trả analytics object cho biểu đồ\nEND FUNCTION",
            },
            {
                "name": "Notebook APIs\nget/add/update/deleteNotebookEntry",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Quản lý sổ tay từ vựng cá nhân.",
                "flow": "FUNCTION notebook\n  GET: query notebook theo userId, phân trang\n  POST: kiểm tra word tồn tại, thêm ghi chú cá nhân\n  PUT: kiểm tra ownership, cập nhật note/favorite\n  DELETE: kiểm tra ownership, xóa entry\n  Trả dữ liệu notebook mới nhất\nEND FUNCTION",
            },
            {
                "name": "Daily goal / Notifications APIs",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Cấu hình mục tiêu ngày, giới hạn SRS và đọc thông báo.",
                "flow": "FUNCTION userSettingsAndNotifications\n  DailyGoal: đọc/cập nhật mục tiêu học mỗi ngày\n  SRSConfig: cập nhật giới hạn số thẻ ôn\n  Notifications: lấy danh sách, mark read, mark all read\n  Trả trạng thái cập nhật cho frontend\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Learning Path & Progress - Lộ trình và thống kê",
        "purpose": "Tạo roadmap học TOEIC theo cấp độ và cung cấp thống kê tiến độ tổng quát.",
        "components": "learning-path.controller.js; learning-path.service.js; progress.routes.js; ProgressController; ProgressService; LearningPathLevels, LearningPathTopics, UserWordProgress.",
        "flow": "Client -> /api/user/learning-path hoặc /api/progress -> Controller -> Service -> SQL Server -> roadmap/progress response.",
        "functions": [
            {
                "name": "GET /learning-path\ngetRoadmap(userId)",
                "file": "controllers/learning-path.controller.js; services/learning-path.service.js",
                "role": "Lấy lộ trình học theo level TOEIC và trạng thái của user.",
                "flow": "FUNCTION getRoadmap\n  Controller lấy userId\n  Service ensureSchema\n  syncPublishedTopics vào roadmap\n  Query levels và topics\n  Lấy progress của user\n  buildRoadmap để gom nhóm\n  Trả roadmap cho frontend\nEND FUNCTION",
            },
            {
                "name": "syncPublishedTopics()",
                "file": "services/learning-path.service.js",
                "role": "Đồng bộ topic đã xuất bản vào bảng lộ trình.",
                "flow": "FUNCTION syncPublishedTopics\n  Query Topics Published\n  FOR mỗi topic\n    Nếu chưa có trong LearningPathTopics\n      INSERT topic vào level phù hợp\n  END FOR\n  Trả số topic đã đồng bộ\nEND FUNCTION",
            },
            {
                "name": "buildRoadmap(levels, topics, progress)",
                "file": "services/learning-path.service.js",
                "role": "Định dạng dữ liệu roadmap cho frontend.",
                "flow": "FUNCTION buildRoadmap\n  FOR mỗi level\n    Lấy topics thuộc level\n    Gắn wordCount, learnedCount, progressPercent\n    Tính trạng thái locked/available/completed\n  END FOR\n  Trả mảng level có topic con\nEND FUNCTION",
            },
            {
                "name": "GET /progress\nProgressService.getProgress",
                "file": "controllers/progress.controller.js; services/progress.service.js",
                "role": "Lấy tiến độ học theo user.",
                "flow": "FUNCTION getProgress\n  Lấy userId từ request\n  Query UserWordProgress\n  Tính số từ đã học, mastered, reviewing\n  Trả progress summary\nEND FUNCTION",
            },
            {
                "name": "GET /progress/stats\nProgressService.getStats",
                "file": "controllers/progress.controller.js; services/progress.service.js",
                "role": "Lấy thống kê tổng quan phục vụ dashboard.",
                "flow": "FUNCTION getStats\n  Query attempts và progress của user\n  Tính accuracy, streak, phân bố mastery\n  Trả stats object\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Gamification - XP, level, achievement",
        "purpose": "Tăng động lực học bằng XP, level, streak và huy hiệu.",
        "components": "GamificationController; GamificationService; user.routes gamification endpoints; bảng UserXPEvents, Achievements, UserAchievements.",
        "flow": "Hoạt động học -> UserService/GamificationController -> GamificationService -> ghi XP events -> cập nhật Users.TotalXP -> kiểm tra achievements -> response.",
        "functions": [
            {
                "name": "GET /gamification/profile\ngetProfile",
                "file": "controllers/gamification.controller.js; services/gamification.service.js",
                "role": "Lấy hồ sơ XP, level, achievements của learner.",
                "flow": "FUNCTION getProfile\n  ensureSchema\n  getMetrics(userId)\n  getAchievements(userId)\n  getLevelState(totalXP)\n  Trả profile gồm XP, level, achievements\nEND FUNCTION",
            },
            {
                "name": "POST /gamification/practice-complete\ncompletePractice",
                "file": "controllers/gamification.controller.js",
                "role": "Cộng XP khi hoàn thành phiên luyện tập.",
                "flow": "FUNCTION completePractice\n  Nhận event practice complete\n  Gọi GamificationService.awardXP\n  Gọi checkAchievements\n  Trả xpGained và achievement mới\nEND FUNCTION",
            },
            {
                "name": "awardXP(userId, event)",
                "file": "services/gamification.service.js",
                "role": "Ghi nhận XP theo loại hoạt động học tập.",
                "flow": "FUNCTION awardXP\n  Xác định XP_REWARDS theo eventType\n  Nếu sourceKey đã tồn tại thì không cộng trùng\n  INSERT UserXPEvents\n  UPDATE Users.TotalXP\n  Tính level mới\n  checkAchievements\n  Trả xpGained, totalXP, levelState\nEND FUNCTION",
            },
            {
                "name": "checkAchievements(userId)",
                "file": "services/gamification.service.js",
                "role": "Mở khóa huy hiệu khi user đạt điều kiện.",
                "flow": "FUNCTION checkAchievements\n  Lấy metrics học tập\n  Lấy achievements chưa đạt\n  FOR mỗi achievement\n    Kiểm tra điều kiện XP/streak/mastery/test\n    Nếu đạt -> INSERT UserAchievements\n  END FOR\n  Trả newlyUnlocked\nEND FUNCTION",
            },
            {
                "name": "PUT /achievements/seen\nmarkAchievementsSeen",
                "file": "controllers/gamification.controller.js; services/gamification.service.js",
                "role": "Đánh dấu huy hiệu đã xem.",
                "flow": "FUNCTION markAchievementsSeen\n  Nhận achievementIds\n  UPDATE UserAchievements SET SeenAt\n  WHERE userId và achievementId hợp lệ\n  Trả success\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Admin - Quản trị hệ thống và nội dung",
        "purpose": "Quản lý toàn bộ dữ liệu hệ thống: topics, words, questions, mini-tests, users, reports, review, analytics, audit logs và notifications.",
        "components": "admin.routes.js; AdminController; AdminService; ReportService; validate/auth middleware; các bảng nội dung, Users, AdminAuditLogs, ContentReviewLogs, Notifications.",
        "flow": "Admin Client -> /api/admin -> verifyToken -> checkPermission/checkAnyPermission -> validate nếu có -> AdminController -> AdminService/ReportService -> SQL Server -> audit log -> response.",
        "functions": [
            {
                "name": "Topic APIs\ngetTopics/createTopic/updateTopic/deleteTopic",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Quản lý chủ đề học TOEIC.",
                "flow": "FUNCTION topicCrud\n  GET: đọc page/limit/filter -> getTopics -> query + count\n  POST: validate body -> createTopic -> INSERT Topics\n  PUT: validate body -> updateTopic -> UPDATE Topics\n  DELETE: deleteTopic -> kiểm tra liên kết -> xóa/ẩn topic\n  Ghi audit log cho thao tác thay đổi\nEND FUNCTION",
            },
            {
                "name": "Topic Category APIs\nget/create/update/deleteTopicCategory",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Quản lý nhóm phân loại chủ đề.",
                "flow": "FUNCTION topicCategoryCrud\n  Kiểm tra quyền MANAGE_TOPIC_CATEGORIES\n  Controller chuyển params/body vào service\n  Service kiểm tra trùng tên/mã\n  INSERT/UPDATE/DELETE TopicCategories\n  Trả danh mục mới nhất hoặc success\nEND FUNCTION",
            },
            {
                "name": "Word APIs\ngetWords/getWordDetail/createWord/updateWord/deleteWord",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Quản lý ngân hàng từ vựng TOEIC.",
                "flow": "FUNCTION wordCrud\n  GET: lọc search/status/topic/partOfSpeech -> query words\n  Detail: lấy word + topics + examples + questions\n  Create/Update: transaction Words + WordTopics + ExampleSentences\n  Delete: archive hoặc hard delete theo quyền\n  Ghi AdminAuditLogs\nEND FUNCTION",
            },
            {
                "name": "Import Words\npreviewWordImport/bulkImportWords",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Nhập hàng loạt từ vựng từ CSV/plain text.",
                "flow": "FUNCTION importWords\n  Preview: parse input -> normalize key -> validate rows -> trả valid/errors\n  Bulk: parse lại input -> transaction\n  FOR mỗi dòng hợp lệ\n    INSERT Words, WordTopics, ExampleSentences\n  END FOR\n  Commit hoặc rollback\n  Trả số dòng thành công/thất bại\nEND FUNCTION",
            },
            {
                "name": "Question APIs\ngetQuestionsByWord\ncreateQuestion\nupdateQuestion\ndeleteQuestion\nbulkImportQuestions",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Quản lý câu hỏi luyện tập gắn với từ vựng.",
                "flow": "FUNCTION questionManagement\n  Lấy wordId/filter từ request\n  Service query Questions theo WordID\n  Create/Update: validate question, answer, options\n  BulkImport: parse CSV -> insert nhiều câu hỏi\n  Delete: xóa câu hỏi hoặc chuyển trạng thái\n  Trả kết quả cho admin UI\nEND FUNCTION",
            },
            {
                "name": "MiniTest APIs\ngetMiniTests\ncreateMiniTest\nupdateMiniTest\ndeleteMiniTest\npublishMiniTest\narchiveMiniTest",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Tạo và quản lý bài kiểm tra ngắn.",
                "flow": "FUNCTION miniTestManagement\n  GET: phân trang/filter MiniTests\n  Create/Update: transaction MiniTests + MiniTestItems\n  Publish/Archive: setMiniTestStatus\n  Delete: kiểm tra attempt/liên kết rồi xóa hoặc archive\n  Ghi audit/content review log\nEND FUNCTION",
            },
            {
                "name": "Student/User APIs\ngetStudents\ncreateUser\nupdateUser\ndeleteUser\ntoggleStudentStatus\nupdateUserRole",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Quản lý tài khoản và phân quyền.",
                "flow": "FUNCTION userManagement\n  GET: lọc user theo role/status/search\n  Create/Update: validate dữ liệu user\n  Toggle: đảo trạng thái IsActive\n  updateUserRole: kiểm tra role hợp lệ -> UPDATE RoleID\n  Delete: xóa mềm hoặc vô hiệu hóa user\nEND FUNCTION",
            },
            {
                "name": "Content Review APIs\ngetPendingContent\nupdateContentStatus\napproveContent\nrejectContent\narchiveContent\ngetContentReviewLogs",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Duyệt nội dung do Creator gửi lên.",
                "flow": "FUNCTION adminContentReview\n  getPendingContent: UNION Topics/Words/Questions/MiniTests PendingReview\n  approve/reject/archive: xác định entityType\n  UPDATE ContentStatus\n  INSERT ContentReviewLogs\n  Nếu reject thì lưu comment/lý do\n  Trả trạng thái mới\nEND FUNCTION",
            },
            {
                "name": "Dashboard & Analytics\ngetStats\ngetAnalytics\ngetContentManagement",
                "file": "controllers/admin.controller.js; services/admin.service.js",
                "role": "Cung cấp KPI và biểu đồ quản trị.",
                "flow": "FUNCTION adminAnalytics\n  Query tổng users/words/questions/topics/minitests\n  Query hoạt động học và attempts\n  Query pending review/content status\n  Gom thành dashboard/analytics object\n  Trả cho frontend admin\nEND FUNCTION",
            },
            {
                "name": "Reports/Audit/Notifications\ngetReports\nupdateReport\ngetAuditLogs\ngetNotifications\nsendAnnouncement\ncreateDailyReminders",
                "file": "controllers/admin.controller.js; services/admin.service.js; services/report.service.js",
                "role": "Xử lý báo cáo, nhật ký quản trị và thông báo hệ thống.",
                "flow": "FUNCTION adminOperations\n  Reports: get/update report status qua ReportService\n  Audit: query AdminAuditLogs theo filter\n  Notifications: query/gửi announcement\n  DailyReminders: tìm user có từ đến hạn ôn -> tạo notification\n  Trả kết quả vận hành\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Creator - Biên tập nội dung học",
        "purpose": "Cho Content Creator tạo topic, word, question, mini-test, media và gửi nội dung sang quy trình duyệt.",
        "components": "creator.routes.js; CreatorController; CreatorService; Review workflow; bảng Topics, Words, Questions, MiniTests, Media, ContentReviewLogs.",
        "flow": "Creator Client -> /api/creator -> verifyToken -> permission middleware -> validate -> CreatorController -> CreatorService -> SQL Server. Nội dung tạo ở Draft, submit-review chuyển PendingReview.",
        "functions": [
            {
                "name": "Dashboard/Analytics\ngetDashboard\ngetContentSummary\ngetTopicAnalytics\ngetMiniTestAnalytics",
                "file": "controllers/creator.controller.js; services/creator.service.js",
                "role": "Cho creator xem số liệu nội dung mình tạo.",
                "flow": "FUNCTION creatorDashboard\n  Lấy userId creator\n  Query số topic/word/question/minitest theo CreatedByUserID\n  Query trạng thái Draft/Pending/Published/Rejected\n  Query analytics theo topic/minitest nếu có id\n  Trả dashboard data\nEND FUNCTION",
            },
            {
                "name": "Creator Topic APIs\ngetTopics\ncreateTopic\nupdateTopic\ndeleteTopic\nsubmitTopicForReview",
                "file": "controllers/creator.controller.js; services/creator.service.js",
                "role": "Creator tạo và gửi duyệt chủ đề.",
                "flow": "FUNCTION creatorTopic\n  GET: lấy topic của creator hoặc Published được phép dùng\n  Create: INSERT Topics status Draft\n  Update/Delete: kiểm tra ownership\n  Submit: submitForReview(Topics, TopicID)\n  Trả topic/status mới\nEND FUNCTION",
            },
            {
                "name": "Creator Word APIs\ngetWords\ncreateWord\nupdateWord\ndeleteWord\nsubmitWordForReview",
                "file": "controllers/creator.controller.js; services/creator.service.js",
                "role": "Creator biên soạn từ vựng và ví dụ.",
                "flow": "FUNCTION creatorWord\n  Kiểm tra PartOfSpeech và topic hợp lệ\n  Create/Update trong transaction\n  INSERT/UPDATE Words, WordTopics, ExampleSentences\n  Delete: kiểm tra ownership\n  Submit: chuyển Draft/Rejected -> PendingReview\nEND FUNCTION",
            },
            {
                "name": "Creator Question APIs\ngetQuestions\ncreateQuestion\nupdateQuestion\ndeleteQuestion\nsubmitQuestionForReview",
                "file": "controllers/creator.controller.js; services/creator.service.js",
                "role": "Creator tạo câu hỏi luyện tập.",
                "flow": "FUNCTION creatorQuestion\n  Kiểm tra word thuộc quyền dùng\n  Validate question type/options/correctAnswer\n  INSERT/UPDATE Questions\n  Delete nếu còn thuộc creator\n  Submit review để admin duyệt\nEND FUNCTION",
            },
            {
                "name": "Creator MiniTest APIs\ngetMiniTests\ncreateMiniTest\nupdateMiniTest\ndeleteMiniTest\naddMiniTestItem\nremoveMiniTestItem\nsubmitMiniTestForReview",
                "file": "controllers/creator.controller.js; services/creator.service.js",
                "role": "Creator tạo bài kiểm tra ngắn từ các câu hỏi.",
                "flow": "FUNCTION creatorMiniTest\n  Validate topicId và questionIds\n  Create/Update MiniTests status Draft\n  addItem/removeItem cập nhật MiniTestItems\n  Submit review chuyển PendingReview\n  Trả mini-test hiện tại\nEND FUNCTION",
            },
            {
                "name": "Media APIs\ngetMedia\ncreateMedia\ndeleteMedia",
                "file": "controllers/creator.controller.js; services/creator.service.js",
                "role": "Quản lý media minh họa nội dung.",
                "flow": "FUNCTION mediaManagement\n  GET: query media của creator\n  POST: validate metadata/url/type -> INSERT MediaAssets\n  DELETE: kiểm tra ownership -> xóa media\n  Trả danh sách hoặc success\nEND FUNCTION",
            },
            {
                "name": "submitForReview(tableName,idColumn,id,userId)",
                "file": "services/creator.service.js",
                "role": "Hàm dùng chung để gửi duyệt nhiều loại entity.",
                "flow": "FUNCTION submitForReview\n  Kiểm tra entity tồn tại và CreatedByUserID=userId\n  Kiểm tra status Draft hoặc Rejected\n  UPDATE ContentStatus = PendingReview\n  UpdatedAt = now\n  Trả entity đã gửi duyệt\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Review - Kiểm duyệt nội dung",
        "purpose": "Tách riêng quy trình duyệt nội dung dùng cho reviewer/admin.",
        "components": "review.routes.js; ReviewController; ReviewService; ContentReviewLogs; Topics/Words/Questions/MiniTests.",
        "flow": "Reviewer -> /api/review -> verifyToken -> REVIEW_CONTENT permission -> ReviewController -> ReviewService -> transaction SQL -> response.",
        "functions": [
            {
                "name": "GET /pending\ngetPendingContent",
                "file": "controllers/review.controller.js; services/review.service.js",
                "role": "Lấy tất cả entity đang PendingReview.",
                "flow": "FUNCTION getPendingContent\n  UNION Topics, Words, Questions, MiniTests\n  JOIN Users để lấy creatorName\n  WHERE ContentStatus='PendingReview'\n  ORDER BY createdAt ASC\n  Trả danh sách pending\nEND FUNCTION",
            },
            {
                "name": "POST /:entityType/:id/approve\napprove",
                "file": "controllers/review.controller.js; services/review.service.js",
                "role": "Duyệt nội dung sang Published.",
                "flow": "FUNCTION approve\n  Resolve entityType -> table/idCol\n  Bắt đầu transaction\n  UPDATE ContentStatus='Published'\n  Set PublishedAt/ReviewedBy/ReviewedAt\n  Nếu MiniTest thì IsPublished=1\n  INSERT ContentReviewLogs\n  Commit và trả true\nEND FUNCTION",
            },
            {
                "name": "POST /reject\nreject",
                "file": "controllers/review.controller.js; services/review.service.js",
                "role": "Từ chối nội dung kèm lý do.",
                "flow": "FUNCTION reject\n  Resolve entity\n  Bắt đầu transaction\n  UPDATE ContentStatus='Rejected'\n  Set ReviewedBy/ReviewedAt\n  INSERT ContentReviewLogs với reason\n  Commit và trả true\nEND FUNCTION",
            },
            {
                "name": "POST /archive\narchive",
                "file": "controllers/review.controller.js; services/review.service.js",
                "role": "Lưu trữ nội dung không còn sử dụng.",
                "flow": "FUNCTION archive\n  Resolve entity\n  UPDATE ContentStatus='Archived'\n  Nếu MiniTest thì IsPublished=0\n  INSERT ContentReviewLogs\n  Trả true nếu có row affected\nEND FUNCTION",
            },
            {
                "name": "GET /logs\ngetReviewLogs",
                "file": "controllers/review.controller.js; services/review.service.js",
                "role": "Xem lịch sử duyệt của một entity.",
                "flow": "FUNCTION getReviewLogs\n  Nhận entityType/entityId\n  Query ContentReviewLogs\n  JOIN Users người duyệt\n  ORDER BY CreatedAt DESC\n  Trả danh sách log\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module AI - Hỗ trợ soạn nội dung TOEIC",
        "purpose": "Gợi ý nghĩa, ví dụ TOEIC, bản dịch và dữ liệu tham khảo để tăng tốc tạo nội dung.",
        "components": "ai.routes.js; AiController; AiService; AI provider; dictionary lookup.",
        "flow": "Admin/Creator -> /api/ai -> verifyToken + permission -> AiController -> AiService -> AI provider/dictionary -> chuẩn hóa JSON -> response.",
        "functions": [
            {
                "name": "POST /suggest-word-content\nsuggestWordContent",
                "file": "controllers/ai.controller.js; services/ai.service.js",
                "role": "Sinh gợi ý nội dung cho một từ vựng TOEIC.",
                "flow": "FUNCTION suggestWordContent\n  Nhận term, meaning, partOfSpeech, exampleCount\n  Làm sạch input\n  Tạo prompt theo ngữ cảnh TOEIC\n  Gọi AI với timeout\n  extractResponseText\n  parseJsonText\n  normalizeExamples\n  Trả suggestion JSON\nEND FUNCTION",
            },
            {
                "name": "generateToeicExamples",
                "file": "services/ai.service.js",
                "role": "Sinh câu ví dụ theo văn phong TOEIC.",
                "flow": "FUNCTION generateToeicExamples\n  Tạo prompt yêu cầu số ví dụ\n  Gọi AI provider\n  Parse JSON/list text\n  Chuẩn hóa sentence + meaning\n  Trả examples\nEND FUNCTION",
            },
            {
                "name": "translateToVietnamese",
                "file": "services/ai.service.js",
                "role": "Dịch nhanh văn bản tiếng Anh sang tiếng Việt.",
                "flow": "FUNCTION translateToVietnamese\n  Làm sạch text\n  Nếu text rỗng -> trả rỗng\n  Gọi AI/dịch vụ dịch\n  Trích xuất response\n  Trả bản dịch tiếng Việt\nEND FUNCTION",
            },
            {
                "name": "lookupDictionary",
                "file": "services/ai.service.js",
                "role": "Tra cứu thông tin từ điển khi cần bổ sung dữ liệu từ.",
                "flow": "FUNCTION lookupDictionary\n  Nhận term\n  Gọi nguồn dictionary\n  Parse phonetic/meaning/examples nếu có\n  Trả dữ liệu tham khảo\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module Report & Notification - Báo cáo và thông báo",
        "purpose": "Người học gửi báo cáo lỗi/nội dung; admin xử lý báo cáo và gửi thông báo học tập.",
        "components": "UserController.createReport; ReportService; AdminController reports/notifications; AdminService notification helpers; Notifications, UserReports.",
        "flow": "Learner tạo report -> ReportService lưu DB. Admin xem/cập nhật report và tạo notification/announcement -> AdminService ghi Notifications -> User đọc thông báo.",
        "functions": [
            {
                "name": "POST /user/report\nReportService.createReport",
                "file": "controllers/user.controller.js; services/report.service.js",
                "role": "Learner báo lỗi hoặc góp ý nội dung.",
                "flow": "FUNCTION createReport\n  Nhận reportData\n  Infer entityType nếu có entityId\n  Validate category/status\n  INSERT Reports/UserReports\n  Trả report vừa tạo\nEND FUNCTION",
            },
            {
                "name": "Admin getReports/updateReport",
                "file": "controllers/admin.controller.js; services/report.service.js",
                "role": "Admin xem và cập nhật trạng thái xử lý báo cáo.",
                "flow": "FUNCTION reportAdmin\n  getReports: phân trang + filter status/type\n  updateReport: validate status/priority/note\n  UPDATE report và reviewer/adminId\n  Trả report mới nhất\nEND FUNCTION",
            },
            {
                "name": "sendAnnouncement",
                "file": "services/admin.service.js",
                "role": "Gửi thông báo chung cho nhóm người dùng.",
                "flow": "FUNCTION sendAnnouncement\n  Nhận audience, title, message\n  Xác định danh sách user nhận\n  INSERT Notifications cho từng user\n  Ghi audit log\n  Trả số thông báo đã tạo\nEND FUNCTION",
            },
            {
                "name": "createDailyReminders",
                "file": "services/admin.service.js",
                "role": "Tạo lời nhắc ôn tập hằng ngày cho user có từ đến hạn.",
                "flow": "FUNCTION createDailyReminders\n  Query learner active\n  Với mỗi learner đếm từ due\n  Nếu dueCount > 0 -> INSERT notification\n  Trả số reminder đã tạo\nEND FUNCTION",
            },
            {
                "name": "User notifications\nget/markRead/markAllRead",
                "file": "controllers/user.controller.js; services/user.service.js",
                "role": "Learner đọc và đánh dấu thông báo.",
                "flow": "FUNCTION userNotifications\n  GET: query Notifications theo userId\n  markRead: UPDATE ReadAt cho notificationId\n  markAllRead: UPDATE tất cả unread của user\n  Trả success/list notifications\nEND FUNCTION",
            },
        ],
    },
    {
        "name": "Module System/Common - Hạ tầng backend",
        "purpose": "Cung cấp kết nối CSDL, validate request, xử lý lỗi, health-check, integration-test và đóng server an toàn.",
        "components": "config/db.js; validate.js; errorHandler.js; health-check.js; integration-test.js; index.js; Express app.",
        "flow": "Request -> Express app -> CORS/JSON parser -> Route -> Middleware -> Controller. Error được chuyển qua next(error) về errorHandler. App dùng poolPromise để tái sử dụng kết nối SQL.",
        "functions": [
            {
                "name": "poolPromise/sql config",
                "file": "config/db.js",
                "role": "Khởi tạo và chia sẻ kết nối SQL Server.",
                "flow": "FUNCTION dbConnection\n  Đọc biến môi trường DB_SERVER/DB_NAME/DB_AUTH\n  Tạo config mssql/msnodesqlv8\n  Khởi tạo connection pool\n  Export poolPromise và sql\nEND FUNCTION",
            },
            {
                "name": "validate(schema)",
                "file": "middlewares/validate.js",
                "role": "Kiểm tra body/query/params trước controller.",
                "flow": "FUNCTION validate\n  Nhận schema\n  Validate req.body theo schema\n  Nếu lỗi -> 400 + error details\n  Nếu hợp lệ -> gán value đã chuẩn hóa\n  next()\nEND FUNCTION",
            },
            {
                "name": "errorHandler(err,req,res,next)",
                "file": "middlewares/errorHandler.js",
                "role": "Chuẩn hóa lỗi toàn backend.",
                "flow": "FUNCTION errorHandler\n  Đọc statusCode/message từ error\n  Nếu production thì ẩn stack\n  Trả JSON { message, details }\n  Kết thúc response\nEND FUNCTION",
            },
            {
                "name": "runHealthCheck",
                "file": "config/health-check.js",
                "role": "Kiểm tra môi trường và kết nối CSDL.",
                "flow": "FUNCTION runHealthCheck\n  Kiểm tra biến môi trường bắt buộc\n  Kết nối poolPromise\n  Query sys.tables\n  In kết quả OK/FAILED\n  process.exit(0)\nEND FUNCTION",
            },
            {
                "name": "runIntegrationTest",
                "file": "config/integration-test.js",
                "role": "Kiểm thử nhanh tích hợp database/schema.",
                "flow": "FUNCTION runIntegrationTest\n  Kết nối DB\n  Chạy các query kiểm tra bảng/dữ liệu mẫu\n  In kết quả từng bước\n  Thoát process theo pass/fail\nEND FUNCTION",
            },
            {
                "name": "gracefulShutdown",
                "file": "index.js",
                "role": "Đóng server an toàn khi nhận signal/lỗi.",
                "flow": "FUNCTION gracefulShutdown\n  Nhận signal\n  server.close để ngừng nhận request mới\n  Đóng kết nối nếu cần\n  process.exit sau khi cleanup\nEND FUNCTION",
            },
        ],
    },
]


def remove_old_class_tables(doc):
    body = doc.element.body
    children = list(body.iterchildren())

    start = None
    end = None
    seen_class_caption = False
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = "".join(child.itertext())
            if "Hình 8:" in text:
                seen_class_caption = True
                continue
            if seen_class_caption and start is None:
                start = i
            if seen_class_caption and "THIẾT KẾ DỮ LIỆU" in text:
                end = i
                break

    if start is None or end is None or start >= end:
        raise RuntimeError("Không tìm thấy vùng bảng thiết kế lớp cần thay thế")

    ref = children[end]
    for child in children[start:end]:
        body.remove(child)
    return ref


def main():
    doc = Document(SRC)
    ref = remove_old_class_tables(doc)

    intro = move_paragraph_before(
        doc,
        ref,
        "Các bảng dưới đây mô tả thiết kế backend theo từng module. Mỗi module nêu rõ thành phần tham gia, luồng chạy tổng quát và luồng xử lý của các hàm/API quan trọng.",
        style=None,
    )
    intro.paragraph_format.space_after = Pt(8)

    for idx, module in enumerate(MODULES, 1):
        move_paragraph_before(doc, ref, f"{idx}. {module['name']}", style="Heading 3")
        table = build_module_table(doc, module)
        move_table_before(doc, ref, table)
        spacer = move_paragraph_before(doc, ref, "", style=None)
        spacer.paragraph_format.space_after = Pt(6)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
