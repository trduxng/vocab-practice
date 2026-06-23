from __future__ import annotations

import json
import math
import re
import textwrap
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\webvocal\vocab-practice")
BACKEND = ROOT / "backend" / "src"
TEMPLATE = ROOT / "artifacts" / "design-doc" / "DetailDesignTemplateV1.0.docx"
PSEUDOCODE_JSON = ROOT / "artifacts" / "design-doc" / "backend_pseudocode.json"
OUTPUT = ROOT / "artifacts" / "design-doc" / "DetailDesignTemplateV1.0_VocaBoost_Backend_FullPseudocode.docx"


DOMAIN_LABELS = {
    "auth": "xác thực",
    "categories": "danh mục nền tảng",
    "admin": "quản trị hệ thống",
    "user": "người học",
    "creator": "biên tập nội dung",
    "ai": "AI hỗ trợ nội dung",
    "progress": "tiến độ học tập",
    "review": "duyệt nội dung",
    "learning-path": "lộ trình học tập",
    "learning_path": "lộ trình học tập",
    "gamification": "gamification và phần thưởng",
    "report": "báo lỗi nội dung",
    "db": "kết nối cơ sở dữ liệu",
    "health-check": "kiểm tra sức khỏe hệ thống",
    "integration-test": "kiểm thử tích hợp",
    "index": "khởi tạo ứng dụng Express",
}


NOUN_MAP = {
    "Topics": "danh sách chủ đề",
    "Topic": "chủ đề",
    "TopicCategories": "danh mục chủ đề",
    "TopicCategory": "danh mục chủ đề",
    "Words": "danh sách từ vựng",
    "Word": "từ vựng",
    "WordDetail": "chi tiết từ vựng",
    "WordImport": "dữ liệu import từ vựng",
    "Questions": "danh sách câu hỏi",
    "Question": "câu hỏi",
    "QuestionsByWord": "các câu hỏi của một từ vựng",
    "MiniTests": "danh sách bài kiểm tra ngắn",
    "MiniTest": "bài kiểm tra ngắn",
    "MiniTestDetails": "chi tiết bài kiểm tra ngắn",
    "MiniTestAnalytics": "thống kê bài kiểm tra ngắn",
    "MiniTestItem": "một phần tử câu hỏi của bài kiểm tra",
    "Stats": "thống kê",
    "DashboardStats": "thống kê dashboard",
    "Analytics": "dữ liệu phân tích",
    "ContentSummary": "tổng hợp trạng thái nội dung",
    "TopicAnalytics": "thống kê theo chủ đề",
    "PartOfSpeeches": "danh mục từ loại",
    "Profile": "hồ sơ người dùng",
    "DueFlashcards": "bộ flashcard đến hạn ôn",
    "Flashcards": "bộ flashcard",
    "TopicWords": "danh sách từ vựng của chủ đề",
    "MasteryTimeline": "dòng thời gian mức độ ghi nhớ",
    "DailyGoal": "mục tiêu hằng ngày",
    "DailyProgress": "tiến độ mục tiêu hằng ngày",
    "SRSConfig": "cấu hình giới hạn ôn SRS",
    "ProgressAnalytics": "phân tích tiến độ học tập",
    "SmartReviewQueue": "hàng đợi ôn tập thông minh",
    "SessionSummary": "tóm tắt phiên ôn tập",
    "MistakeReviewQueue": "hàng đợi ôn lại câu sai",
    "Notebook": "sổ tay từ vựng",
    "NotebookEntry": "một mục trong sổ tay từ vựng",
    "Notifications": "danh sách thông báo",
    "Notification": "thông báo",
    "Reports": "danh sách báo lỗi/nội dung",
    "Report": "báo lỗi/nội dung",
    "PendingContent": "danh sách nội dung chờ duyệt",
    "ReviewLogs": "nhật ký duyệt nội dung",
    "AuditLogs": "nhật ký audit quản trị",
    "Students": "danh sách người dùng/học viên",
    "StudentStatus": "trạng thái tài khoản học viên",
    "UserRole": "vai trò người dùng",
    "ContentStatus": "trạng thái nội dung",
    "OverviewStats": "thống kê tổng quan",
    "WeeklyActivity": "hoạt động theo tuần",
    "TodayActivity": "hoạt động trong ngày",
    "RecentUsers": "danh sách người dùng gần đây",
    "TopCourses": "danh sách chủ đề/bộ học nổi bật",
    "Dashboard": "dashboard",
    "Roadmap": "lộ trình học tập",
    "Pending": "danh sách chờ xử lý",
    "Media": "danh sách media metadata",
    "MediaAsset": "tài nguyên media",
    "AchievementsSeen": "trạng thái đã xem huy hiệu",
}


SPECIFIC_METHOD_DESCRIPTIONS = {
    ("services/admin.service.js", "createWord"): "Tạo mới từ vựng trong transaction, gán chủ đề, thêm ví dụ minh họa và ghi nhật ký quản trị.",
    ("services/admin.service.js", "updateWord"): "Cập nhật từ vựng, làm mới liên kết chủ đề/câu ví dụ trong transaction và lưu audit log.",
    ("services/admin.service.js", "deleteWord"): "Xóa mềm hoặc xóa logic từ vựng theo ràng buộc nghiệp vụ và cập nhật nhật ký thao tác.",
    ("services/admin.service.js", "bulkInsertWords"): "Đọc dữ liệu import hàng loạt, ánh xạ topic/từ loại và tạo từng từ vựng vào CSDL.",
    ("services/admin.service.js", "previewWordImport"): "Tiền kiểm dữ liệu CSV import từ vựng, ánh xạ reference data và trả danh sách lỗi theo từng dòng.",
    ("services/user.service.js", "getDueFlashcards"): "Lấy danh sách flashcard đến hạn ôn theo SRS, có thể lọc theo chủ đề hoặc chế độ new/learned.",
    ("services/user.service.js", "submitAnswer"): "Ghi nhận câu trả lời, gọi stored procedure lưu attempt, cập nhật lịch SRS theo reviewRating và cộng XP nếu cần.",
    ("services/user.service.js", "submitWordReview"): "Cập nhật trực tiếp UserWordProgress theo kết quả ôn từ, tính mastery/ease factor/next review date và gamification.",
    ("services/user.service.js", "getSmartReviewQueue"): "Xây dựng hàng đợi ôn thông minh dựa trên độ quá hạn, số lần sai liên tiếp và mức độ thành thạo.",
    ("services/auth.service.js", "login"): "Xác thực email/mật khẩu, tải quyền từ RolePermissions, sinh JWT và cộng điểm đăng nhập hằng ngày cho learner.",
    ("services/auth.service.js", "register"): "Kiểm tra trùng email, băm mật khẩu, gán role Learner mặc định và tạo tài khoản mới.",
    ("services/review.service.js", "approve"): "Phê duyệt nội dung chờ duyệt, cập nhật trạng thái Published, ghi mốc reviewed/published và lưu review log.",
    ("services/review.service.js", "reject"): "Từ chối nội dung PendingReview, lưu lý do từ chối và ghi nhận lịch sử duyệt.",
    ("services/review.service.js", "archive"): "Chuyển nội dung sang Archived và ghi nhận thao tác lưu trữ vào bảng log duyệt.",
    ("services/ai.service.js", "suggestWordContent"): "Điều phối tra từ điển, dịch nghĩa và gọi OpenAI để gợi ý nội dung từ vựng TOEIC hoàn chỉnh.",
    ("middlewares/auth.js", "verifyToken"): "Kiểm tra Bearer token JWT, giải mã payload và gắn thông tin người dùng vào req.",
    ("middlewares/auth.js", "checkPermission"): "Tạo middleware kiểm tra một quyền cụ thể trước khi cho phép truy cập endpoint.",
    ("middlewares/auth.js", "checkAnyPermission"): "Tạo middleware cho phép truy cập nếu người dùng có ít nhất một quyền trong danh sách.",
    ("middlewares/auth.js", "checkOwnership"): "Kiểm tra quyền sở hữu bản ghi khi người dùng không có đặc quyền quản trị hệ thống.",
    ("config/db.js", "poolPromise"): "Khởi tạo và tái sử dụng connection pool SQL Server cho toàn bộ backend.",
}


ROUTE_SPECIFIC_DESCRIPTIONS = {
    ("POST", "/submit-answer"): "Nhận kết quả trả lời bài tập/flashcard từ người học, cập nhật attempt và logic SRS.",
    ("GET", "/flashcards"): "Trả về bộ flashcard hiện tại cho learner dựa trên tiến độ và lịch ôn.",
    ("PUT", "/goals/srs-config"): "Cập nhật giới hạn số lượng mục ôn SRS trong mỗi phiên học.",
    ("GET", "/review/smart-queue"): "Trả danh sách ôn ưu tiên theo điểm nguy cơ quên.",
    ("POST", "/words/import-preview"): "Phân tích CSV import từ vựng và trả về kết quả kiểm tra trước khi ghi dữ liệu.",
    ("POST", "/words/bulk-import"): "Thực thi import hàng loạt từ vựng vào cơ sở dữ liệu.",
}


CORE_ALGORITHMS = [
    {
        "title": "Đăng nhập người dùng (AuthService.login)",
        "objective": "Xác thực tài khoản, tải quyền, phát hành JWT và kích hoạt phần thưởng đăng nhập hằng ngày.",
        "inputs": ["email", "password"],
        "outputs": ["token JWT", "thông tin người dùng", "danh sách quyền", "gamification nếu là learner"],
        "steps": [
            "Truy vấn người dùng theo email và trạng thái IsActive = 1.",
            "So khớp mật khẩu nhập vào với PasswordHash bằng bcrypt.compare.",
            "Nếu sai thông tin đăng nhập thì ném lỗi nghiệp vụ và kết thúc.",
            "Truy vấn RolePermissions + Permissions để lấy toàn bộ permission code của người dùng.",
            "Tạo payload JWT gồm id, fullName, role, permissions; ký bằng JWT_SECRET trong 1 ngày.",
            "Nếu role là Learner thì thử cộng thưởng daily login thông qua GamificationService.",
            "Trả về token, thông tin người dùng và phần thưởng (nếu có).",
        ],
        "pseudocode": """Hàm login(email, password):\n  user <- query Users + Roles theo email và IsActive = 1\n  nếu user rỗng: ném lỗi sai tài khoản/mật khẩu\n  nếu bcrypt.compare(password, user.passwordHash) = false: ném lỗi\n  permissions <- query RolePermissions/Permissions theo user.id\n  payload <- {id, fullName, role, permissions}\n  token <- jwt.sign(payload, JWT_SECRET, expiresIn = 1 day)\n  nếu user.role = Learner:\n      gamification <- awardDailyLogin(user.id)\n  trả về {token, user, permissions, gamification}""",
    },
    {
        "title": "Lấy flashcard đến hạn ôn theo SRS (UserService.getDueFlashcards)",
        "objective": "Xây dựng bộ câu hỏi/flashcard ưu tiên theo lịch ôn, chủ đề và trạng thái học tập của learner.",
        "inputs": ["userId", "topicId (optional)", "mode = new | learned | null"],
        "outputs": ["Danh sách flashcard đã trộn và kèm metadata SRS"],
        "steps": [
            "Đọc SRSReviewLimit của người dùng; nếu chưa có thì dùng mặc định 15.",
            "Nếu client chưa truyền topicId, hệ thống lấy chủ đề người học đang enrol đầu tiên còn active/published.",
            "Ghép dữ liệu Words với PartOfSpeeches, Questions, ExampleSentences và UserWordProgress.",
            "Lọc theo ContentStatus = Published và theo topic hoặc trạng thái đã từng học.",
            "Nếu mode = new thì chỉ lấy các từ chưa có repetition count; nếu learned thì chỉ lấy từ đã có progress; nếu null thì lấy từ chưa đến hạn hoặc đã đến hạn tùy luật SRS.",
            "Sắp xếp ưu tiên: mục quá hạn trước, rồi theo NextReviewDate, MasteryLevel và random để giảm lặp nhàm chán.",
            "Trả về tối đa @Limit bản ghi dùng cho phiên học hiện tại.",
        ],
        "pseudocode": """Hàm getDueFlashcards(userId, topicId, mode):\n  limit <- Users.SRSReviewLimit hoặc 15\n  nếu topicId rỗng:\n      topicId <- chủ đề learner đang theo học đầu tiên còn hiệu lực\n  query Words + Questions + Examples + UserWordProgress\n  lọc Published và lọc theo topic phù hợp\n  nếu mode = 'new': chỉ lấy repetitionCount = 0\n  nếu mode = 'learned': chỉ lấy bản ghi đã có UserWordProgress\n  nếu mode khác: chỉ lấy từ chưa có lịch hoặc NextReviewDate <= now\n  order by overdue first, next review, mastery, random\n  trả về TOP(limit)""",
    },
    {
        "title": "Gửi kết quả trả lời và cập nhật lịch ôn (UserService.submitAnswer)",
        "objective": "Lưu lịch sử làm bài, điều chỉnh EaseFactor/NextReviewDate theo rating Again-Hard-Good-Easy và cộng XP khi cần.",
        "inputs": ["userId", "questionId", "submittedAnswer", "reviewRating", "activityType"],
        "outputs": ["Kết quả attempt", "feedback SRS", "gamification/xpGained"],
        "steps": [
            "Gọi stored procedure usp_SubmitQuestionAttempt để lưu ExerciseAttempts và lấy WordID chuẩn của câu hỏi.",
            "Nếu WordID không ánh xạ được thì trả lỗi vì câu hỏi không gắn với từ vựng hợp lệ.",
            "Nếu reviewRating thuộc Again/Hard/Good/Easy thì cập nhật EaseFactor và NextReviewDate trong UserWordProgress.",
            "Again giảm EaseFactor mạnh và dời lịch 10 phút; Hard dời 1 ngày; Good/Easy dời 1-30 ngày theo MasteryLevel.",
            "Nếu activityType = LearnWord thì gọi GamificationService.awardXP để cộng điểm.",
            "Trộn dữ liệu attempt + feedback + gamification thành object trả về cho frontend.",
        ],
        "pseudocode": """Hàm submitAnswer(payload):\n  result <- execute usp_SubmitQuestionAttempt(userId, questionId, submittedAnswer)\n  wordId <- result.WordID\n  nếu wordId rỗng: ném lỗi\n  nếu reviewRating hợp lệ:\n      update UserWordProgress\n         EaseFactor <- tăng/giảm theo rating\n         NextReviewDate <- 10 phút / 1 ngày / 1-30 ngày theo rating + mastery\n  nếu activityType = LearnWord:\n      gamification <- awardXP(...)\n  trả về merge(result, reviewFeedback, gamification)""",
    },
    {
        "title": "Cập nhật trực tiếp tiến độ từ vựng (UserService.submitWordReview)",
        "objective": "Tạo mới hoặc cập nhật UserWordProgress cho learner sau mỗi lần ôn từ độc lập ngoài bài tập.",
        "inputs": ["userId", "wordId", "isCorrect", "reviewRating", "activityType"],
        "outputs": ["masteryLevel", "memoryStatus", "nextReviewDate", "xpGained"],
        "steps": [
            "Dùng MERGE trên UserWordProgress để xử lý cả trường hợp đã có bản ghi và chưa có bản ghi.",
            "Khi bản ghi tồn tại: tăng/giảm MasteryLevel, cộng RepetitionCount, cập nhật chuỗi đúng/sai liên tiếp.",
            "Tính NextReviewDate theo reviewRating hoặc theo việc đúng/sai với các mốc 10 phút, 1 ngày, 3-30 ngày.",
            "Điều chỉnh EaseFactor trong khoảng tối thiểu/tối đa để phản ánh độ khó của từ với người học.",
            "Gán MemoryStatus thành Learning / Reviewing / Mastered / Lapsed tùy mastery và kết quả hiện tại.",
            "Nếu chưa có bản ghi thì insert bản ghi mặc định với EaseFactor = 2.50 và lịch ôn ban đầu.",
            "Nếu activityType = LearnWord thì cộng XP và trả toàn bộ feedback cho frontend.",
        ],
        "pseudocode": """Hàm submitWordReview(userId, wordId, isCorrect, reviewRating):\n  MERGE UserWordProgress theo (UserID, WordID)\n    nếu MATCHED:\n      cập nhật mastery, repetition, consecutiveCorrect/Wrong\n      tính NextReviewDate theo rating + mastery\n      tính EaseFactor theo Again/Hard/Good/Easy\n      cập nhật MemoryStatus\n    nếu NOT MATCHED:\n      insert bản ghi khởi tạo với mastery ban đầu\n  nếu activityType = LearnWord: cộng XP\n  trả về feedback SRS mới nhất""",
    },
    {
        "title": "Tạo từ vựng mới ở khu vực quản trị (AdminService.createWord)",
        "objective": "Ghi một từ vựng mới cùng liên kết chủ đề và câu ví dụ, bảo đảm tính toàn vẹn bằng transaction.",
        "inputs": ["wordData: term, meaning, phonetic, partOfSpeechId, topicIds, examples, status", "adminId"],
        "outputs": ["id", "term", "meaning"],
        "steps": [
            "Kiểm tra status nội dung hợp lệ và chuẩn hóa danh sách topicIds, loại bỏ giá trị trùng/rỗng.",
            "Lọc danh sách examples chỉ giữ các ví dụ có sentence không rỗng.",
            "Mở transaction SQL Server để bảo đảm các insert xảy ra đồng bộ.",
            "Insert bản ghi mới vào bảng Words và lấy WordID vừa sinh.",
            "Lặp qua từng topicId để thêm bản ghi WordTopics.",
            "Lặp qua từng ví dụ hợp lệ để thêm ExampleSentences.",
            "Commit transaction, ghi AdminAuditLog và trả object kết quả.",
            "Nếu xảy ra lỗi ở bất kỳ bước nào thì rollback transaction và ném lỗi ra ngoài controller.",
        ],
        "pseudocode": """Hàm createWord(wordData, adminId):\n  validate status\n  topicIds <- unique(topicIds hợp lệ)\n  examples <- lọc câu ví dụ có nội dung\n  begin transaction\n  thử:\n      wordId <- insert Words\n      cho mỗi topicId: insert WordTopics(wordId, topicId)\n      cho mỗi example: insert ExampleSentences(wordId, ...)\n      commit\n      logAdminAction('CREATE_WORD')\n      trả về {id: wordId, term, meaning}\n  nếu lỗi:\n      rollback\n      throw error""",
    },
    {
        "title": "Tiền kiểm import CSV từ vựng (AdminService.previewWordImport)",
        "objective": "Phân tích dữ liệu import và báo lỗi theo từng dòng trước khi ghi thật vào CSDL.",
        "inputs": ["input CSV/text/json rows"],
        "outputs": ["total", "valid", "invalid", "rows[] có cờ hợp lệ/lỗi"],
        "steps": [
            "Parse input thành danh sách rows, tự nhận biết delimiter và header aliases.",
            "Tải reference data gồm PartOfSpeeches và Topics từ cơ sở dữ liệu.",
            "Tạo map tra cứu theo id, name và code đã normalize để hỗ trợ nhập CSV linh hoạt.",
            "Duyệt từng dòng, đọc các cột term/meaning/phonetic/partOfSpeech/topic/example qua danh sách alias.",
            "Kiểm tra dữ liệu bắt buộc, tra cứu partOfSpeech, ánh xạ topic theo id hoặc tên/mã chủ đề.",
            "Thu thập errors theo từng dòng; tăng bộ đếm valid/invalid tương ứng.",
            "Trả previewRows cho frontend để hiển thị trước khi gọi bulk-import thật sự.",
        ],
        "pseudocode": """Hàm previewWordImport(input):\n  rows <- parseWordImport(input)\n  partsOfSpeech, topics <- query reference data\n  build map theo id + tên + code\n  cho từng row:\n      term, meaning, phonetic <- lấy theo aliases\n      validate partOfSpeech và topics\n      nếu lỗi: mark invalid + lưu errors\n      ngược lại: mark valid\n  trả về {total, valid, invalid, rows}""",
    },
    {
        "title": "Duyệt nội dung chờ duyệt (ReviewService.approve / reject / archive)",
        "objective": "Quản lý vòng đời Topic/Word/Question/MiniTest ở trạng thái PendingReview bằng transaction và log lịch sử.",
        "inputs": ["entityType", "entityId", "adminId", "reason (đối với reject)"],
        "outputs": ["true/false tùy có cập nhật được hay không"],
        "steps": [
            "Gọi _resolveEntity để ánh xạ entityType thành tên bảng, khóa chính và các cột reviewed/published tương ứng.",
            "Bắt đầu transaction để đồng bộ cập nhật trạng thái và log lịch sử duyệt.",
            "Approve: đổi ContentStatus thành Published; riêng MiniTest còn bật IsPublished = 1.",
            "Reject: đổi ContentStatus thành Rejected và lưu reason vào ContentReviewLogs.",
            "Archive: đổi ContentStatus thành Archived, lưu comment mặc định vào log.",
            "Nếu không có dòng nào bị cập nhật thì rollback và trả false.",
            "Nếu thành công thì insert một dòng vào ContentReviewLogs rồi commit transaction.",
        ],
        "pseudocode": """Hàm approve(entityType, entityId, adminId):\n  entityMeta <- resolveEntity(entityType)\n  begin transaction\n  update bảng đích từ PendingReview -> Published\n  nếu entityType = MiniTest: set IsPublished = 1\n  insert ContentReviewLogs\n  commit hoặc rollback nếu rowsAffected = 0\n\nHàm reject(..., reason):\n  update PendingReview -> Rejected\n  insert log kèm reason\n\nHàm archive(...):\n  update bất kỳ trạng thái -> Archived\n  insert log archive""",
    },
]


def iter_backend_files() -> list[Path]:
    files = [BACKEND / "index.js"]
    for folder in ["config", "middlewares", "routes", "controllers", "services"]:
        files.extend(sorted((BACKEND / folder).glob("*.js")))
    return files


def humanize_slug(slug: str) -> str:
    return DOMAIN_LABELS.get(slug, slug.replace("-", " ").replace("_", " "))


def prettify_filename(path: Path) -> str:
    stem = path.stem
    if stem == "index":
        return "ApplicationServer"
    parts = stem.split(".")
    if len(parts) > 1:
        stem = parts[0]
    text = re.sub(r"[-_]+", " ", stem).title().replace("Ai", "AI").replace("Db", "DB")
    suffix = path.parent.name[:-1].title() if path.parent.name.endswith("s") else path.parent.name.title()
    if suffix in {"Controller", "Service", "Route"}:
        return text.replace(" ", "") + suffix
    if path.parent.name == "middlewares":
        return text.replace(" ", "") + "Middleware"
    if path.parent.name == "config":
        return text.replace(" ", "") + "Config"
    return text.replace(" ", "")


def remove_body_preserve_section(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(font_size)
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_run_font(run, *, size=12, bold=False, italic=False, color=None, name="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.font.name = name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def add_paragraph(doc: Document, text: str = "", *, style=None, align=None, size=12, bold=False, italic=False, color=None, space_after=6, space_before=0):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 3)}"
    run = p.add_run(text)
    size = {1: 16, 2: 14, 3: 12}.get(level, 12)
    color = {1: "2F5496", 2: "2F5496", 3: "1F1F1F"}.get(level, "000000")
    set_run_font(run, size=size, bold=True, color=color)
    p.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc: Document, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    return table


def add_table_title(doc: Document, title: str):
    p = add_paragraph(doc, title, size=11, bold=True, space_after=3, space_before=6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_simple_kv_table(doc: Document, title: str, rows: list[tuple[str, str]], widths=(2.0, 4.5)):
    add_table_title(doc, title)
    table = add_table(doc, len(rows) + 1, 2)
    headers = ["Thuộc tính", "Mô tả"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, (k, v) in enumerate(rows, start=1):
        set_cell_text(table.rows[r].cells[0], k, bold=True)
        set_cell_text(table.rows[r].cells[1], v)
    set_table_widths(table, widths)
    return table


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def camel_words(name: str) -> list[str]:
    return re.findall(r"[A-Z]?[a-z]+|[A-Z]+(?=[A-Z]|$)|\d+", name)


def object_phrase(name: str) -> str:
    if name in NOUN_MAP:
        return NOUN_MAP[name]
    words = camel_words(name)
    joined = "".join(words)
    if joined in NOUN_MAP:
        return NOUN_MAP[joined]
    phrase = " ".join(word.lower() for word in words)
    return phrase or name


def describe_by_action(action: str, obj: str, folder: str) -> str:
    layer_prefix = {
        "controllers": "Nhận request HTTP, kiểm tra/chuẩn hóa tham số và gọi service để ",
        "services": "Xử lý nghiệp vụ và truy vấn/cập nhật CSDL để ",
        "middlewares": "Thực thi middleware để ",
        "config": "Cấu hình tiện ích hệ thống nhằm ",
    }.get(folder, "Thực hiện xử lý để ")

    base = {
        "get": f"lấy {obj}",
        "create": f"tạo mới {obj}",
        "update": f"cập nhật {obj}",
        "delete": f"xóa {obj}",
        "remove": f"gỡ bỏ {obj}",
        "add": f"thêm {obj}",
        "submit": f"gửi {obj}",
        "preview": f"xem trước {obj}",
        "bulkImport": f"nhập hàng loạt {obj}",
        "bulkInsert": f"ghi hàng loạt {obj}",
        "approve": f"phê duyệt {obj}",
        "reject": f"từ chối {obj}",
        "archive": f"lưu trữ {obj}",
        "publish": f"xuất bản {obj}",
        "toggle": f"bật/tắt {obj}",
        "mark": f"đánh dấu {obj}",
        "check": f"kiểm tra {obj}",
        "ensure": f"bảo đảm sẵn sàng {obj}",
        "parse": f"phân tích và chuẩn hóa {obj}",
        "normalize": f"chuẩn hóa {obj}",
        "build": f"xây dựng {obj}",
        "infer": f"suy luận {obj}",
        "split": f"tách {obj}",
        "award": f"cộng thưởng cho {obj}",
        "sync": f"đồng bộ {obj}",
        "register": f"đăng ký {obj}",
        "login": f"xác thực {obj}",
        "logout": f"kết thúc phiên của {obj}",
        "with": f"xử lý {obj}",
    }
    return layer_prefix + base.get(action, f"xử lý {obj}") + "."


def generic_method_description(file_key: str, method: str, folder: str) -> str:
    specific = SPECIFIC_METHOD_DESCRIPTIONS.get((file_key, method))
    if specific:
        return specific

    prefixes = [
        "bulkImport", "bulkInsert", "get", "create", "update", "delete", "remove",
        "add", "submit", "preview", "approve", "reject", "archive", "publish",
        "toggle", "mark", "check", "ensure", "parse", "normalize", "build",
        "infer", "split", "award", "sync", "register", "login", "logout",
    ]
    for prefix in prefixes:
        if method.startswith(prefix):
            obj = object_phrase(method[len(prefix):] or prefix)
            return describe_by_action(prefix, obj, folder)

    if method.startswith("_"):
        return f"Hàm nội bộ dùng để hỗ trợ lớp/module trong việc {object_phrase(method.lstrip('_'))}."
    return f"Hàm hỗ trợ xử lý {object_phrase(method)} trong module backend."


def file_domain(path: Path) -> str:
    stem = path.stem
    if "." in stem:
        stem = stem.split(".")[0]
    return stem


def module_responsibility(path: Path) -> str:
    folder = path.parent.name
    domain = humanize_slug(file_domain(path))
    if path.name == "index.js":
        return "Khởi tạo máy chủ Express, gắn middleware toàn cục, mount route và xử lý vòng đời tiến trình."
    mapping = {
        "routes": f"Khai báo các endpoint HTTP cho phân hệ {domain}, gắn middleware quyền và nối sang controller.",
        "controllers": f"Nhận request/response cho phân hệ {domain}, gọi service và chuẩn hóa HTTP status trả về.",
        "services": f"Thực hiện toàn bộ nghiệp vụ chính của phân hệ {domain}, làm việc trực tiếp với SQL Server và stored procedure khi cần.",
        "middlewares": f"Cung cấp middleware cắt ngang cho {domain}, tập trung vào xác thực, phân quyền, validate và chuẩn hóa lỗi.",
        "config": f"Cấu hình và tiện ích hệ thống phục vụ {domain}.",
    }
    return mapping.get(folder, f"Module backend cho {domain}.")


def parse_functions(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    names = []
    patterns = [
        r"static\s+async\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(",
        r"static\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(",
        r"const\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*(?:async\s*)?\(",
        r"function\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(",
        r"async\s+function\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(",
    ]
    for pattern in patterns:
        for match in re.findall(pattern, text):
            if match not in names and match not in {"if", "for", "while", "switch", "catch", "return"}:
                names.append(match)

    if path.name == "validate.js":
        extra = ["validate", "schemas.register", "schemas.login", "schemas.createWord", "schemas.createQuestion", "schemas.miniTest", "schemas.createReport"]
        for item in extra:
            if item not in names:
                names.append(item)

    if path.name == "db.js":
        for item in ["config", "poolPromise", "sql"]:
            if item not in names:
                names.append(item)
    return names


def parse_routes(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    compact = re.sub(r"\s+", " ", text)
    routes = []
    pattern = re.compile(
        r"router\.(get|post|put|delete|patch)\(\s*['\"]([^'\"]+)['\"]\s*,\s*(.*?)\);"
    )
    for verb, route_path, rest in pattern.findall(compact):
        middlewares = []
        for token in ["verifyToken", "checkPermission", "checkAnyPermission", "validate", "express.text"]:
            if token in rest:
                middlewares.append(token)
        controller_match = re.search(r"([A-Za-z]+Controller)\.([A-Za-z_][A-Za-z0-9_]*)", rest)
        handler = controller_match.group(2) if controller_match else ""
        routes.append({
            "verb": verb.upper(),
            "path": route_path,
            "middlewares": middlewares,
            "handler": handler,
        })
    return routes


def describe_route(route: dict, parent_path: str = "") -> str:
    specific = ROUTE_SPECIFIC_DESCRIPTIONS.get((route["verb"], route["path"]))
    if specific:
        return specific
    if route["handler"]:
        handler_phrase = object_phrase(route["handler"])
        text = f"Endpoint {route['verb']} {parent_path}{route['path']} dùng để xử lý {handler_phrase}"
    else:
        text = f"Endpoint {route['verb']} {parent_path}{route['path']} của backend"
    if route["middlewares"]:
        text += f", có gắn middleware {', '.join(route['middlewares'])}"
    return text + "."


def module_inventory():
    modules = []
    for path in iter_backend_files():
        rel = path.relative_to(BACKEND)
        if path.parent.name == "routes":
            modules.append({
                "type": "route",
                "path": path,
                "rel": str(rel).replace("\\", "/"),
                "name": prettify_filename(path),
                "domain": file_domain(path),
                "responsibility": module_responsibility(path),
                "properties": [
                    ("Tầng", "Route layer"),
                    ("File", str(rel).replace("\\", "/")),
                    ("Trách nhiệm", module_responsibility(path)),
                ],
                "routes": parse_routes(path),
            })
        else:
            modules.append({
                "type": "module",
                "path": path,
                "rel": str(rel).replace("\\", "/"),
                "name": prettify_filename(path),
                "domain": file_domain(path),
                "responsibility": module_responsibility(path),
                "properties": [
                    ("Tầng", path.parent.name[:-1].title() if path.parent.name.endswith("s") else path.parent.name.title()),
                    ("File", str(rel).replace("\\", "/")),
                    ("Trách nhiệm", module_responsibility(path)),
                ],
                "methods": parse_functions(path),
            })
    return modules


def set_page_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def configure_headers(doc: Document):
    for idx, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        for para in header.paragraphs:
            clear_paragraph(para)
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("VOCAB-PRACTICE - Thiết kế chi tiết backend")
        set_run_font(run, size=11, color="666666")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        first_header = section.first_page_header
        if not first_header.paragraphs:
            first_header.add_paragraph()
        for para in first_header.paragraphs:
            clear_paragraph(para)


def add_cover(doc: Document):
    add_paragraph(doc, "WEBSITE LUYỆN TẬP TỪ VỰNG THI TOEIC (VOCABOOST)", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, color="1F1F1F", space_after=10, space_before=40)
    add_paragraph(doc, "THIẾT KẾ CHI TIẾT BACKEND", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, color="2F5496", space_after=18)
    add_paragraph(doc, "Phiên bản hiệu chỉnh từ template DetailDesignTemplateV1.0", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True, color="555555", space_after=18)

    meta_rows = [
        ("Mã dự án", "VOCAB-PRACTICE"),
        ("Tên dự án", "TOEIC Vocabulary Learning Platform / VocaBoost"),
        ("Phạm vi tài liệu", "Thiết kế chi tiết backend, API, middleware, service và luồng xử lý nghiệp vụ"),
        ("Phiên bản", "V3.0 - Backend Detailed Design + Full Pseudocode"),
        ("Ngày", "23/06/2026"),
    ]
    add_simple_kv_table(doc, "Thông tin tài liệu", meta_rows)

    add_paragraph(doc, "", space_after=12)
    sign_table = add_table(doc, 4, 3)
    set_table_widths(sign_table, (1.8, 2.8, 1.8))
    headers = ["Vai trò", "Người phụ trách", "Ngày"]
    for idx, header in enumerate(headers):
        set_cell_text(sign_table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, row in enumerate([
        ("Người lập", "Codex / nhóm phát triển", "23/06/2026"),
        ("Người kiểm tra", "Nhóm dự án", ""),
        ("Người phê duyệt", "Chủ nhiệm đề tài", ""),
    ], start=1):
        for col, value in enumerate(row):
            set_cell_text(sign_table.rows[row_idx].cells[col], value, align=WD_ALIGN_PARAGRAPH.CENTER if col == 2 else WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def add_toc(doc: Document):
    add_paragraph(
        doc,
        "MỤC LỤC",
        size=16,
        bold=True,
        color="2F5496",
        space_before=14,
        space_after=8,
    )
    items = [
        "1. Giới thiệu",
        "2. Tổng quan hệ thống",
        "3. Kiến trúc hệ thống backend",
        "4. Thiết kế lớp và module backend",
        "5. Giả mã chi tiết từng hàm backend",
        "6. Thuật toán nghiệp vụ trọng yếu",
        "7. Thiết kế dữ liệu backend",
        "8. Liên kết frontend - backend",
        "9. Kết luận và khuyến nghị mở rộng",
    ]
    for item in items:
        add_paragraph(doc, item, size=12, space_after=2)
    doc.add_page_break()


def add_intro(doc: Document):
    add_heading(doc, "GIỚI THIỆU", 1)
    add_heading(doc, "Mục đích", 2)
    add_paragraph(
        doc,
        "Tài liệu này mô tả chi tiết phần backend của hệ thống VocaBoost - nền tảng học từ vựng TOEIC có áp dụng Spaced Repetition System (SRS). "
        "Mục tiêu của tài liệu là giúp nhóm phát triển, kiểm thử và giảng viên/nhà hướng dẫn hiểu rõ cấu trúc mã nguồn, trách nhiệm của từng module, "
        "các hàm quan trọng, các endpoint API và các luồng xử lý nghiệp vụ cốt lõi.",
        size=12,
    )
    add_heading(doc, "Phạm vi", 2)
    add_paragraph(
        doc,
        "Phạm vi tài liệu tập trung vào backend Node.js/Express trong thư mục backend/src, bao gồm: cấu hình hệ thống, middleware, route, controller, service, "
        "các thuật toán SRS, quản lý nội dung, gamification, lộ trình học và các API hỗ trợ AI. Phần frontend chỉ được nhắc đến ở mức liên kết API tiêu biểu.",
        size=12,
    )
    add_heading(doc, "Vai trò và trách nhiệm", 2)
    for line in [
        "Nhà phát triển backend: tra cứu trách nhiệm từng file, từng service và luồng nghiệp vụ để bảo trì/mở rộng hệ thống.",
        "Nhóm kiểm thử: dựa trên route/controller/service để xây dựng test case cho API và kiểm thử tích hợp.",
        "Giảng viên hoặc người phản biện đồ án: theo dõi đầy đủ cách hệ thống hiện thực yêu cầu phân quyền, quản trị nội dung và SRS.",
        "Nhóm frontend: đối chiếu các API được backend công bố và chuẩn dữ liệu trả về.",
    ]:
        add_paragraph(doc, f"• {line}", size=12, space_after=2)
    add_heading(doc, "Tài liệu tham khảo", 2)
    refs = [
        ("README.md", "Mô tả tổng quan dự án, stack công nghệ, chức năng lõi và hướng dẫn chạy."),
        ("Database/prototype_database.sql", "Script tạo CSDL nền ban đầu của dự án."),
        ("Database/TOEIC_Schema_dbdiagram.dbml", "Sơ đồ logic và các bảng chính dùng cho backend."),
        ("Database/giaithichbang.md", "Giải thích nghiệp vụ và tác dụng của các bảng dữ liệu."),
    ]
    add_simple_kv_table(doc, "Bảng tài liệu tham khảo", refs)
    add_heading(doc, "Từ và thuật ngữ", 2)
    terms = add_table(doc, 6, 3)
    set_table_widths(terms, (1.5, 1.5, 3.5))
    for idx, header in enumerate(["Thuật ngữ", "Viết tắt", "Diễn giải"]):
        set_cell_text(terms.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    glossary = [
        ("SRS", "Spaced Repetition System", "Cơ chế lặp lại ngắt quãng dùng để xác định thời điểm ôn từ vựng tối ưu."),
        ("JWT", "JSON Web Token", "Token xác thực mà backend phát hành sau khi đăng nhập thành công."),
        ("CRUD", "Create/Read/Update/Delete", "Bốn thao tác cơ bản trên dữ liệu quản trị nội dung."),
        ("Learner", "Người học", "Vai trò sử dụng hệ thống để học và kiểm tra từ vựng."),
        ("Content Creator", "Biên tập viên", "Vai trò tạo chủ đề, từ vựng, câu hỏi, mini test và gửi duyệt."),
    ]
    for row_idx, row in enumerate(glossary, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(terms.rows[row_idx].cells[col_idx], value)


def add_system_overview(doc: Document):
    add_heading(doc, "TỔNG QUAN HỆ THỐNG", 1)
    add_paragraph(
        doc,
        "VocaBoost là website học từ vựng TOEIC theo mô hình full-stack. Frontend Next.js gọi backend Express thông qua các API REST dưới tiền tố /api. "
        "Backend đảm nhiệm xác thực, phân quyền, điều phối logic SRS, quản lý nội dung, quản lý báo lỗi, duyệt nội dung, gamification và hỗ trợ AI sinh ví dụ/định nghĩa.",
        size=12,
    )
    add_heading(doc, "Các phân hệ chức năng chính", 2)
    subsystems = [
        ("Phân hệ người học", "Học flashcard, mini test, dashboard tiến độ, notebook, notifications, daily goal và hàng đợi SRS."),
        ("Phân hệ biên tập viên", "Tạo chủ đề/từ vựng/câu hỏi/mini test, quản lý media và gửi nội dung đi duyệt."),
        ("Phân hệ quản trị viên", "Duyệt nội dung, quản trị người dùng, báo cáo, thông báo, thống kê và import dữ liệu."),
        ("Phân hệ AI", "Gợi ý nghĩa, phiên âm và câu ví dụ TOEIC bằng OpenAI + Dictionary API + Google Translate."),
    ]
    add_simple_kv_table(doc, "Bảng phân hệ", subsystems)
    add_heading(doc, "Luồng request backend tổng quát", 2)
    flow_rows = [
        ("B1", "Frontend gọi apiClient tới endpoint /api/... tương ứng."),
        ("B2", "Express nhận request tại index.js và định tuyến sang route module phù hợp."),
        ("B3", "Middleware verifyToken/checkPermission/validate được áp dụng trước khi vào controller."),
        ("B4", "Controller tách tham số, gọi service, ánh xạ lỗi nghiệp vụ sang HTTP status."),
        ("B5", "Service truy vấn SQL Server, chạy transaction/stored procedure và gom business result."),
        ("B6", "Controller trả JSON về frontend; errorHandler xử lý lỗi còn sót ở cuối pipeline."),
    ]
    add_simple_kv_table(doc, "Luồng backend chuẩn", flow_rows, widths=(0.8, 5.7))


def add_architecture(doc: Document, modules):
    add_heading(doc, "KIẾN TRÚC HỆ THỐNG BACKEND", 1)
    add_paragraph(
        doc,
        "Backend sử dụng kiến trúc phân lớp mỏng: Route -> Middleware -> Controller -> Service -> SQL Server. "
        "Mỗi lớp có trách nhiệm rõ ràng, giúp dễ mở rộng chức năng và dễ kiểm thử tích hợp.",
        size=12,
    )
    counts = defaultdict(int)
    total_routes = 0
    for module in modules:
        if module["type"] == "route":
            counts["route files"] += 1
            total_routes += len(module["routes"])
        else:
            counts[module["path"].parent.name] += 1
    arch_rows = [
        ("Máy chủ ứng dụng", "Express.js khởi tạo tại backend/src/index.js, mount toàn bộ route dưới /api."),
        ("Tầng Route", f"{counts['route files']} file route, khai báo khoảng {total_routes} endpoint REST."),
        ("Tầng Controller", f"{counts['controllers']} controller chịu trách nhiệm request/response."),
        ("Tầng Service", f"{counts['services']} service xử lý nghiệp vụ và SQL trực tiếp."),
        ("Middleware dùng chung", f"{counts['middlewares']} middleware cho auth, permission, validation, error handling."),
        ("Config/utility", f"{counts['config']} module cấu hình/kết nối/health check."),
    ]
    add_simple_kv_table(doc, "Phân rã backend", arch_rows)
    add_heading(doc, "Thành phần chính", 2)
    component_rows = [
        ("ApplicationServer", "Khởi tạo Express, cấu hình CORS/Helmet/JSON parser, route mount, health endpoint và graceful shutdown."),
        ("Auth Middleware", "Giải mã JWT, kiểm tra quyền, kiểm tra quyền sở hữu nội dung."),
        ("Validate Middleware", "Dùng Zod schemas để chặn dữ liệu request không hợp lệ trước khi vào controller."),
        ("Route Modules", "Định nghĩa đầy đủ endpoint cho auth, learner, admin, creator, AI, progress và review."),
        ("Controller Modules", "Nhận request, chọn status code, chuyển dữ liệu sang service, ánh xạ lỗi nghiệp vụ sang response."),
        ("Service Modules", "Chứa toàn bộ logic truy vấn SQL, transaction, stored procedure, SRS, import, gamification, report và AI."),
        ("SQL Server", "Lưu người dùng, từ vựng, câu hỏi, mini test, tiến độ học, logs, notifications và cấu hình gamification."),
    ]
    add_simple_kv_table(doc, "Mô tả component", component_rows)


def add_route_table(doc: Document, module: dict, base_path: str):
    add_heading(doc, module["name"], 3)
    add_paragraph(doc, module["responsibility"], size=11)
    props = module["properties"] + [("Base path", base_path)]
    add_simple_kv_table(doc, "Thuộc tính module", props)
    add_table_title(doc, "Danh sách endpoint")
    table = add_table(doc, len(module["routes"]) + 1, 4)
    set_table_widths(table, (0.95, 2.1, 1.55, 1.9))
    for idx, header in enumerate(["HTTP", "Path", "Handler", "Mô tả"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, route in enumerate(module["routes"], start=1):
        set_cell_text(table.rows[row_idx].cells[0], route["verb"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9.5)
        set_cell_text(table.rows[row_idx].cells[1], base_path + route["path"], font_size=9.5)
        set_cell_text(table.rows[row_idx].cells[2], route["handler"] or "N/A", font_size=9.5)
        set_cell_text(table.rows[row_idx].cells[3], describe_route(route, base_path), font_size=9.5)


def add_module_table(doc: Document, module: dict):
    add_heading(doc, module["name"], 3)
    add_paragraph(doc, module["responsibility"], size=11)
    add_simple_kv_table(doc, "Thuộc tính module", module["properties"])
    add_table_title(doc, "Danh sách hàm")
    methods = module["methods"] or ["N/A"]
    table = add_table(doc, len(methods) + 1, 3)
    set_table_widths(table, (0.9, 2.2, 3.0))
    for idx, header in enumerate(["Loại", "Hàm / thành phần", "Mô tả"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, method in enumerate(methods, start=1):
        kind = "Method"
        if method.startswith("schemas."):
            kind = "Schema"
        elif method in {"config", "poolPromise", "sql"}:
            kind = "Property"
        elif method.startswith("_"):
            kind = "Private"
        set_cell_text(table.rows[row_idx].cells[0], kind, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[row_idx].cells[1], method)
        set_cell_text(table.rows[row_idx].cells[2], generic_method_description(module["rel"], method, module["path"].parent.name))


def add_module_design(doc: Document, modules):
    add_heading(doc, "THIẾT KẾ LỚP VÀ MODULE BACKEND", 1)
    add_paragraph(
        doc,
        "Mục này mô tả đầy đủ các file backend theo đúng code hiện tại của dự án. Mỗi module được trình bày theo phong cách tương tự tài liệu mẫu: "
        "nêu vai trò, thuộc tính cấu hình quan trọng và danh sách hàm/endpoint cùng chức năng.",
        size=12,
    )

    order = [
        ("Application", lambda m: m["rel"] == "index.js"),
        ("Config", lambda m: m["path"].parent.name == "config"),
        ("Middleware", lambda m: m["path"].parent.name == "middlewares"),
        ("Routes", lambda m: m["path"].parent.name == "routes"),
        ("Controllers", lambda m: m["path"].parent.name == "controllers"),
        ("Services", lambda m: m["path"].parent.name == "services"),
    ]
    route_prefix = {
        "auth.routes.js": "/api/auth",
        "categories.routes.js": "/api/categories",
        "admin.routes.js": "/api/admin",
        "user.routes.js": "/api/user",
        "progress.routes.js": "/api/progress",
        "creator.routes.js": "/api/creator",
        "review.routes.js": "/api/admin/content-review",
        "ai.routes.js": "/api/ai",
    }

    for section_name, predicate in order:
        add_heading(doc, section_name, 2)
        section_modules = [m for m in modules if predicate(m)]
        for module in section_modules:
            if module["type"] == "route":
                add_route_table(doc, module, route_prefix.get(module["path"].name, ""))
            else:
                add_module_table(doc, module)


def load_pseudocode_inventory():
    if not PSEUDOCODE_JSON.exists():
        raise FileNotFoundError(
            f"Chưa có dữ liệu giả mã: {PSEUDOCODE_JSON}. "
            "Hãy chạy extract_backend_pseudocode.js trước khi dựng tài liệu."
        )
    return json.loads(PSEUDOCODE_JSON.read_text(encoding="utf-8"))


def wrap_pseudocode(text: str, width: int = 92) -> list[str]:
    output = []
    for raw_line in text.splitlines():
        stripped = raw_line.lstrip(" ")
        leading = len(raw_line) - len(stripped)
        if not stripped:
            output.append("")
            continue
        subsequent = " " * (leading + 4)
        wrapped = textwrap.wrap(
            stripped,
            width=max(36, width - leading),
            initial_indent=" " * leading,
            subsequent_indent=subsequent,
            break_long_words=False,
            break_on_hyphens=False,
            replace_whitespace=False,
        )
        output.extend(wrapped or [" " * leading])
    return output


def add_pseudocode_box(doc: Document, pseudocode: str):
    table = add_table(doc, 1, 1)
    table.autofit = False
    set_table_widths(table, (6.5,))
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    wrapped_lines = wrap_pseudocode(pseudocode)
    for idx, line in enumerate(wrapped_lines):
        run = p.add_run(line)
        set_run_font(run, size=8.5, name="Courier New", color="1F1F1F")
        if idx < len(wrapped_lines) - 1:
            run.add_break()
    return table


def function_output_description(file_path: str) -> str:
    folder = Path(file_path).parent.name
    if folder == "controllers":
        return "HTTP status và JSON response; lỗi được chuyển tới middleware errorHandler."
    if folder == "middlewares":
        return "Quyết định gọi next(), trả lỗi HTTP hoặc tạo middleware kiểm tra request."
    if folder == "services":
        return "Dữ liệu nghiệp vụ, kết quả cập nhật CSDL hoặc lỗi nghiệp vụ."
    if folder == "config":
        return "Kết quả kiểm tra/cấu hình phục vụ vận hành backend."
    return "Kết quả xử lý tương ứng với từng nhánh RETURN của hàm."


def add_detailed_pseudocode_section(doc: Document, pseudocode_inventory):
    add_heading(doc, "GIẢ MÃ CHI TIẾT TỪNG HÀM BACKEND", 1)
    total_functions = sum(len(file_item.get("functions", [])) for file_item in pseudocode_inventory)
    add_paragraph(
        doc,
        f"Chương này mô tả giả mã cho toàn bộ {total_functions} hàm có tên được phát hiện trong backend/src. "
        "Giả mã được sinh từ cấu trúc câu lệnh thực tế của mã nguồn và giữ lại đầy đủ các nhánh TRY/CATCH, "
        "IF/ELSE, vòng lặp, lời gọi service, transaction, truy vấn SQL, stored procedure, RETURN và HTTP response. "
        "Cách trình bày thống nhất theo mẫu FUNCTION ... END FUNCTION mà nhóm sử dụng trong báo cáo.",
        size=12,
    )

    summary_rows = []
    for file_item in pseudocode_inventory:
        functions = file_item.get("functions", [])
        if functions:
            summary_rows.append((file_item["file"], f"{len(functions)} hàm"))
    add_simple_kv_table(doc, "Thống kê phạm vi giả mã", summary_rows)

    layer_order = ["config", "middlewares", "controllers", "services", "."]
    layer_labels = {
        "config": "Config và kiểm tra vận hành",
        "middlewares": "Middleware",
        "controllers": "Controller",
        "services": "Service",
        ".": "Application bootstrap",
    }
    grouped = defaultdict(list)
    for file_item in pseudocode_inventory:
        parent = Path(file_item["file"]).parent.as_posix()
        grouped[parent].append(file_item)

    function_index = 0
    for layer in layer_order:
        file_items = grouped.get(layer, [])
        if not file_items:
            continue
        add_heading(doc, layer_labels[layer], 2)
        for file_item in file_items:
            functions = file_item.get("functions", [])
            if not functions:
                continue
            add_heading(doc, file_item["file"], 3)
            add_paragraph(
                doc,
                f"File này có {len(functions)} hàm được mô tả. Thứ tự bên dưới bám theo vị trí xuất hiện trong mã nguồn.",
                size=10.5,
                italic=True,
                color="555555",
                space_after=4,
            )
            for function in functions:
                function_index += 1
                signature = function.get("signature") or f"{function['name']}({', '.join(function.get('params', []))})"
                title = add_paragraph(
                    doc,
                    f"Hàm {function_index}: {signature}",
                    size=11.5,
                    bold=True,
                    color="2F5496",
                    space_before=8,
                    space_after=2,
                )
                title.paragraph_format.keep_with_next = True
                description = generic_method_description(
                    file_item["file"],
                    function["name"],
                    Path(file_item["file"]).parent.name,
                )
                metadata = (
                    f"Dòng nguồn: {function.get('startLine', 'N/A')} | "
                    f"Đầu vào: {', '.join(function.get('params', [])) or 'Không có tham số'} | "
                    f"Đầu ra: {function_output_description(file_item['file'])}"
                )
                meta_p = add_paragraph(doc, metadata, size=9.5, italic=True, color="555555", space_after=2)
                meta_p.paragraph_format.keep_with_next = True
                desc_p = add_paragraph(doc, f"Mục đích: {description}", size=10.5, space_after=3)
                desc_p.paragraph_format.keep_with_next = True
                add_pseudocode_box(doc, function["pseudocode"])


def add_algorithm_section(doc: Document):
    add_heading(doc, "THUẬT TOÁN NGHIỆP VỤ TRỌNG YẾU", 1)
    add_paragraph(
        doc,
        "Ngoài danh sách hàm, phần này đi sâu vào những nghiệp vụ cốt lõi nhất của backend để làm rõ cách hệ thống thật sự hoạt động. "
        "Đây là phần quan trọng nếu báo cáo cần thể hiện ‘mô tả chi tiết cách hiện thực’ chứ không chỉ liệt kê tên hàm.",
        size=12,
    )
    for item in CORE_ALGORITHMS:
        add_heading(doc, item["title"], 2)
        add_paragraph(doc, f"Mục tiêu: {item['objective']}", size=11, bold=False, space_after=4)
        add_simple_kv_table(doc, "Đầu vào / đầu ra", [
            ("Đầu vào", "; ".join(item["inputs"])),
            ("Đầu ra", "; ".join(item["outputs"])),
        ])
        add_table_title(doc, "Luồng xử lý")
        table = add_table(doc, len(item["steps"]) + 1, 2)
        set_table_widths(table, (0.6, 5.9))
        set_cell_text(table.rows[0].cells[0], "Bước", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[0].cells[1], "Mô tả", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for idx, step in enumerate(item["steps"], start=1):
            set_cell_text(table.rows[idx].cells[0], f"{idx}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(table.rows[idx].cells[1], step)
        add_table_title(doc, "Giả mã")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(8)
        for line in item["pseudocode"].splitlines():
            run = p.add_run(line + "\n")
            set_run_font(run, size=10.5, name="Courier New")


def add_data_design(doc: Document):
    add_heading(doc, "THIẾT KẾ DỮ LIỆU BACKEND", 1)
    add_paragraph(
        doc,
        "Backend hiện tại làm việc chủ yếu trên SQL Server với các bảng lõi phục vụ người dùng, nội dung, học tập, SRS, mini test, duyệt nội dung và gamification. "
        "Bảng dưới đây tổng hợp những thực thể quan trọng nhất liên quan trực tiếp tới mã nguồn backend.",
        size=12,
    )
    rows = [
        ("Users", "Lưu tài khoản, vai trò, trạng thái hoạt động và cấu hình học tập của người dùng."),
        ("Topics / TopicCategories", "Lưu nhóm chủ đề và các chủ đề học từ vựng TOEIC."),
        ("Words", "Lưu từ vựng, nghĩa, phiên âm, media và trạng thái nội dung."),
        ("ExampleSentences", "Lưu ví dụ câu cho từng từ vựng."),
        ("WordTopics", "Bảng nối nhiều-nhiều giữa Words và Topics."),
        ("Questions", "Lưu câu hỏi trắc nghiệm/fill blank/dictation gắn với từng từ."),
        ("MiniTests / MiniTestItems", "Lưu bài kiểm tra ngắn và danh sách câu hỏi của từng bài."),
        ("UserWordProgress", "Lưu tiến độ nhớ từ, mastery, ease factor, next review date và memory status cho SRS."),
        ("ExerciseAttempts", "Lưu lịch sử làm bài/câu hỏi để tính accuracy và phân tích tiến độ."),
        ("UserTopicEnrollments", "Lưu các chủ đề mà learner đã đăng ký hoặc đang theo học."),
        ("ContentReviewLogs", "Lưu lịch sử duyệt nội dung bởi admin/reviewer."),
        ("ContentReports", "Lưu các báo lỗi, phản hồi nội dung do learner gửi."),
        ("UserXPEvents / UserAchievements", "Lưu sự kiện cộng điểm, huy hiệu và trạng thái gamification."),
        ("Notifications", "Lưu thông báo gửi tới learner hoặc do admin tạo."),
    ]
    add_simple_kv_table(doc, "Các bảng lõi", rows)
    add_heading(doc, "Ràng buộc dữ liệu liên quan đến SRS", 2)
    srs_rows = [
        ("Khóa duy nhất", "UserWordProgress có cặp duy nhất (UserID, WordID) để mỗi người học chỉ có một bản ghi tiến độ trên mỗi từ."),
        ("Chỉ mục chính", "Có index trên (UserID, NextReviewDate) giúp lấy danh sách ôn đến hạn nhanh hơn."),
        ("Thuộc tính quan trọng", "MasteryLevel, EaseFactor, RepetitionCount, ConsecutiveCorrect, ConsecutiveWrong, NextReviewDate, MemoryStatus."),
        ("Ý nghĩa nghiệp vụ", "Cho phép backend lập lịch ôn thông minh, ưu tiên từ yếu và giữ lịch sử tiến bộ lâu dài."),
    ]
    add_simple_kv_table(doc, "Dữ liệu hỗ trợ SRS", srs_rows)


def add_frontend_linkage(doc: Document):
    add_heading(doc, "LIÊN KẾT FRONTEND - BACKEND", 1)
    add_paragraph(
        doc,
        "Frontend Next.js gọi backend chủ yếu thông qua frontend/src/lib/api-client.ts và các service tương ứng. "
        "Bảng dưới đây nêu một số liên kết tiêu biểu để chứng minh backend đã được nối thực tế vào giao diện người dùng.",
        size=12,
    )
    rows = [
        ("auth.service.ts -> /api/auth", "Đăng ký, đăng nhập, lưu token và cookie phục vụ middleware server-side."),
        ("user.service.ts -> /api/user", "Phiên học flashcard, mini test, dashboard, notebook, notifications và mục tiêu hằng ngày."),
        ("admin.service.ts -> /api/admin", "Quản trị chủ đề, từ vựng, câu hỏi, bài kiểm tra, người dùng, báo cáo, thống kê."),
        ("creator.service.ts -> /api/creator", "Luồng biên tập nội dung, media, gửi duyệt nội dung và analytics cá nhân."),
        ("ai.service.ts -> /api/ai", "Gọi backend AI để gợi ý nội dung từ vựng trước khi lưu vào CSDL."),
    ]
    add_simple_kv_table(doc, "Liên kết service frontend", rows)


def add_conclusion(doc: Document):
    add_heading(doc, "KẾT LUẬN VÀ KHUYẾN NGHỊ MỞ RỘNG", 1)
    add_paragraph(
        doc,
        "Phần backend của VocaBoost đã được tổ chức khá rõ theo kiến trúc route-controller-service và đã hiện thực được nhiều yêu cầu quan trọng: "
        "xác thực JWT, phân quyền theo Role/Permission, quản trị nội dung, quy trình duyệt, mini test, báo lỗi, thông báo, gamification và đặc biệt là cơ chế SRS.",
        size=12,
    )
    for line in [
        "Bổ sung test tự động cho các luồng transaction quan trọng như createWord, review approve/reject và submitAnswer.",
        "Tách bớt các service lớn (đặc biệt admin.service.js và user.service.js) thành nhiều service nhỏ để tăng tính maintainable.",
        "Chuẩn hóa thêm logging và observability cho các API AI/upstream để dễ truy vết lỗi môi trường.",
        "Bổ sung tài liệu validate schema và contract API response nếu nhóm muốn hướng tới bàn giao sản phẩm hoàn chỉnh hơn.",
    ]:
        add_paragraph(doc, f"• {line}", size=12, space_after=2)


def build():
    doc = Document(str(TEMPLATE))
    remove_body_preserve_section(doc)
    set_page_defaults(doc)
    configure_headers(doc)
    modules = module_inventory()
    pseudocode_inventory = load_pseudocode_inventory()

    add_cover(doc)
    add_toc(doc)
    add_intro(doc)
    add_system_overview(doc)
    add_architecture(doc, modules)
    add_module_design(doc, modules)
    add_detailed_pseudocode_section(doc, pseudocode_inventory)
    add_algorithm_section(doc)
    add_data_design(doc)
    add_frontend_linkage(doc)
    add_conclusion(doc)

    doc.save(str(OUTPUT))
    print(str(OUTPUT))


if __name__ == "__main__":
    build()
